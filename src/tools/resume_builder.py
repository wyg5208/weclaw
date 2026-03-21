"""简历生成工具 — 专业简历生成与导出。

支持5种简历模板:
- minimal（简约）: 黑白简洁，适合互联网行业
- business（商务）: 深蓝配色，适合金融/管理
- creative（创意）: 彩色侧边栏，适合设计/创意行业
- academic（学术）: 传统学术格式，适合教育/科研
- technical（技术）: 技能进度条展示，适合IT/工程
"""

from __future__ import annotations

import html
import logging
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor

from src.tools.base import ActionDef, BaseTool, ToolResult, ToolResultStatus

logger = logging.getLogger(__name__)


# 模板配置
TEMPLATE_CONFIGS = {
    "minimal": {
        "name": "minimal",
        "title": "简约模板",
        "style": "黑白简洁，无多余装饰",
        "suitable_for": "互联网、科技、初创公司",
        "colors": {
            "primary": RGBColor(0, 0, 0),
            "secondary": RGBColor(80, 80, 80),
            "accent": RGBColor(100, 100, 100),
            "background": None,
        },
        "fonts": {
            "name": "微软雅黑",
            "name_size": 24,
            "section_size": 14,
            "body_size": 11,
        },
    },
    "business": {
        "name": "business",
        "title": "商务模板",
        "style": "深蓝配色，专业稳重",
        "suitable_for": "金融、管理、咨询、银行",
        "colors": {
            "primary": RGBColor(25, 55, 95),
            "secondary": RGBColor(50, 80, 120),
            "accent": RGBColor(0, 102, 204),
            "background": None,
        },
        "fonts": {
            "name": "微软雅黑",
            "name_size": 22,
            "section_size": 13,
            "body_size": 11,
        },
    },
    "creative": {
        "name": "creative",
        "title": "创意模板",
        "style": "彩色侧边栏，活泼现代",
        "suitable_for": "设计、创意、广告、媒体",
        "colors": {
            "primary": RGBColor(255, 87, 51),
            "secondary": RGBColor(52, 73, 94),
            "accent": RGBColor(46, 204, 113),
            "background": RGBColor(245, 245, 245),
        },
        "fonts": {
            "name": "微软雅黑",
            "name_size": 26,
            "section_size": 14,
            "body_size": 11,
        },
    },
    "academic": {
        "name": "academic",
        "title": "学术模板",
        "style": "传统学术格式，正式规范",
        "suitable_for": "教育、科研、学术机构",
        "colors": {
            "primary": RGBColor(0, 0, 0),
            "secondary": RGBColor(60, 60, 60),
            "accent": RGBColor(128, 0, 0),
            "background": None,
        },
        "fonts": {
            "name": "宋体",
            "name_size": 18,
            "section_size": 14,
            "body_size": 12,
        },
    },
    "technical": {
        "name": "technical",
        "title": "技术模板",
        "style": "技能进度条展示，突出技术能力",
        "suitable_for": "IT、工程、软件开发",
        "colors": {
            "primary": RGBColor(39, 174, 96),
            "secondary": RGBColor(44, 62, 80),
            "accent": RGBColor(52, 152, 219),
            "background": RGBColor(236, 240, 241),
        },
        "fonts": {
            "name": "Consolas",
            "name_size": 22,
            "section_size": 13,
            "body_size": 11,
        },
    },
}


class ResumeBuilderTool(BaseTool):
    """专业简历生成工具。

    支持多种模板风格，可生成 DOCX/PDF/HTML 格式简历。
    """

    name = "resume_builder"
    emoji = "📋"
    title = "简历生成器"
    description = "专业简历生成工具"
    timeout = 120

    def __init__(self, output_dir: str = None) -> None:
        """初始化简历生成工具。

        Args:
            output_dir: 输出目录，默认为项目的 generated/日期/ 目录
        """
        super().__init__()
        self.output_dir = (
            Path(output_dir)
            if output_dir
            else Path(__file__).parent.parent.parent
            / "generated"
            / datetime.now().strftime("%Y-%m-%d")
        )
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def get_actions(self) -> list[ActionDef]:
        return [
            ActionDef(
                name="generate_resume",
                description="根据个人信息生成专业简历文档（DOCX格式）",
                parameters={
                    "template": {
                        "type": "string",
                        "description": "简历模板风格",
                        "enum": ["minimal", "business", "creative", "academic", "technical"],
                    },
                    "personal_info": {
                        "type": "object",
                        "description": "个人基本信息",
                        "properties": {
                            "name": {"type": "string", "description": "姓名"},
                            "phone": {"type": "string", "description": "电话"},
                            "email": {"type": "string", "description": "邮箱"},
                            "address": {"type": "string", "description": "地址"},
                        },
                        "required": ["name"],
                    },
                    "education": {
                        "type": "array",
                        "description": "教育经历列表",
                        "items": {
                            "type": "object",
                            "properties": {
                                "school": {"type": "string", "description": "学校名称"},
                                "degree": {"type": "string", "description": "学位"},
                                "major": {"type": "string", "description": "专业"},
                                "start_date": {"type": "string", "description": "开始日期"},
                                "end_date": {"type": "string", "description": "结束日期"},
                                "gpa": {"type": "string", "description": "GPA"},
                            },
                        },
                    },
                    "experience": {
                        "type": "array",
                        "description": "工作经历列表",
                        "items": {
                            "type": "object",
                            "properties": {
                                "company": {"type": "string", "description": "公司名称"},
                                "title": {"type": "string", "description": "职位"},
                                "start_date": {"type": "string", "description": "开始日期"},
                                "end_date": {"type": "string", "description": "结束日期"},
                                "description": {"type": "string", "description": "工作描述"},
                            },
                        },
                    },
                    "skills": {
                        "type": "array",
                        "description": "技能列表",
                        "items": {"type": "string"},
                    },
                    "projects": {
                        "type": "array",
                        "description": "项目经历列表（可选）",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string", "description": "项目名称"},
                                "role": {"type": "string", "description": "担任角色"},
                                "description": {"type": "string", "description": "项目描述"},
                                "technologies": {"type": "string", "description": "使用技术"},
                            },
                        },
                    },
                    "awards": {
                        "type": "array",
                        "description": "获奖荣誉列表（可选）",
                        "items": {"type": "string"},
                    },
                    "summary": {
                        "type": "string",
                        "description": "个人简介/求职意向（可选）",
                    },
                },
                required_params=["template", "personal_info", "education", "experience", "skills"],
            ),
            ActionDef(
                name="list_templates",
                description="列出所有可用的简历模板及其详细信息",
                parameters={},
                required_params=[],
            ),
            ActionDef(
                name="export_resume",
                description="将已生成的DOCX简历导出为其他格式（PDF/HTML）",
                parameters={
                    "input_file": {
                        "type": "string",
                        "description": "输入的DOCX文件路径",
                    },
                    "format": {
                        "type": "string",
                        "description": "导出格式",
                        "enum": ["pdf", "docx", "html"],
                    },
                },
                required_params=["input_file", "format"],
            ),
        ]

    async def execute(self, action: str, params: dict[str, Any]) -> ToolResult:
        """执行指定动作。"""
        if action == "generate_resume":
            return self._generate_resume(params)
        elif action == "list_templates":
            return self._list_templates()
        elif action == "export_resume":
            return self._export_resume(params)
        else:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的动作: {action}",
            )

    def _list_templates(self) -> ToolResult:
        """列出所有可用模板。"""
        templates_info = []
        for key, config in TEMPLATE_CONFIGS.items():
            templates_info.append(
                f"**{config['title']}** (`{key}`)\n"
                f"  - 风格: {config['style']}\n"
                f"  - 适用场景: {config['suitable_for']}"
            )

        output = "📋 可用简历模板:\n\n" + "\n\n".join(templates_info)

        return ToolResult(
            status=ToolResultStatus.SUCCESS,
            output=output,
            data={
                "templates": [
                    {
                        "name": k,
                        "title": v["title"],
                        "style": v["style"],
                        "suitable_for": v["suitable_for"],
                    }
                    for k, v in TEMPLATE_CONFIGS.items()
                ]
            },
        )

    def _generate_resume(self, params: dict[str, Any]) -> ToolResult:
        """生成简历文档。"""
        template = params.get("template", "minimal")
        personal_info = params.get("personal_info", {})
        education = params.get("education", [])
        experience = params.get("experience", [])
        skills = params.get("skills", [])
        projects = params.get("projects", [])
        awards = params.get("awards", [])
        summary = params.get("summary", "")

        if template not in TEMPLATE_CONFIGS:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的模板: {template}，可选: {', '.join(TEMPLATE_CONFIGS.keys())}",
            )

        if not personal_info.get("name"):
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="个人信息中必须包含姓名 (name)",
            )

        config = TEMPLATE_CONFIGS[template]

        try:
            # 创建文档
            doc = Document()

            # 根据模板类型生成不同风格
            if template == "minimal":
                self._build_minimal_resume(doc, config, personal_info, summary, education, experience, skills, projects, awards)
            elif template == "business":
                self._build_business_resume(doc, config, personal_info, summary, education, experience, skills, projects, awards)
            elif template == "creative":
                self._build_creative_resume(doc, config, personal_info, summary, education, experience, skills, projects, awards)
            elif template == "academic":
                self._build_academic_resume(doc, config, personal_info, summary, education, experience, skills, projects, awards)
            elif template == "technical":
                self._build_technical_resume(doc, config, personal_info, summary, education, experience, skills, projects, awards)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name_part = personal_info.get("name", "resume").replace(" ", "_")
            filename = f"resume_{name_part}_{template}_{timestamp}.docx"
            output_path = self.output_dir / filename

            # 保存文档
            doc.save(str(output_path))
            file_size = output_path.stat().st_size

            logger.info("简历生成成功: %s", output_path)

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=(
                    f"✅ 简历已生成\n"
                    f"📁 文件: {output_path.name}\n"
                    f"📊 大小: {file_size} 字节\n"
                    f"🎨 模板: {config['title']}"
                ),
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "template": template,
                    "person_name": personal_info.get("name"),
                },
            )
        except Exception as e:
            logger.error("简历生成失败: %s", e, exc_info=True)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"简历生成失败: {e}",
            )

    def _build_minimal_resume(
        self,
        doc: Document,
        config: dict,
        personal_info: dict,
        summary: str,
        education: list,
        experience: list,
        skills: list,
        projects: list,
        awards: list,
    ) -> None:
        """构建简约风格简历。"""
        colors = config["colors"]
        fonts = config["fonts"]

        # 姓名（居中，大字）
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.font.size = Pt(fonts["name_size"])
        name_run.font.bold = True
        name_run.font.color.rgb = colors["primary"]

        # 联系方式（居中，一行）
        contact_parts = []
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("address"):
            contact_parts.append(personal_info["address"])

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(fonts["body_size"])
            contact_run.font.color.rgb = colors["secondary"]

        # 分隔线
        self._add_horizontal_line(doc)

        # 个人简介
        if summary:
            self._add_section_header(doc, "个人简介", colors["primary"], fonts["section_size"])
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(fonts["body_size"])

        # 工作经历
        if experience:
            self._add_section_header(doc, "工作经历", colors["primary"], fonts["section_size"])
            for exp in experience:
                self._add_experience_item(doc, exp, fonts, colors)

        # 教育经历
        if education:
            self._add_section_header(doc, "教育经历", colors["primary"], fonts["section_size"])
            for edu in education:
                self._add_education_item(doc, edu, fonts, colors)

        # 技能
        if skills:
            self._add_section_header(doc, "专业技能", colors["primary"], fonts["section_size"])
            skills_para = doc.add_paragraph(" • ".join(skills))
            for run in skills_para.runs:
                run.font.size = Pt(fonts["body_size"])

        # 项目经历
        if projects:
            self._add_section_header(doc, "项目经历", colors["primary"], fonts["section_size"])
            for proj in projects:
                self._add_project_item(doc, proj, fonts, colors)

        # 获奖荣誉
        if awards:
            self._add_section_header(doc, "获奖荣誉", colors["primary"], fonts["section_size"])
            for award in awards:
                p = doc.add_paragraph(f"• {award}")
                for run in p.runs:
                    run.font.size = Pt(fonts["body_size"])

    def _build_business_resume(
        self,
        doc: Document,
        config: dict,
        personal_info: dict,
        summary: str,
        education: list,
        experience: list,
        skills: list,
        projects: list,
        awards: list,
    ) -> None:
        """构建商务风格简历。"""
        colors = config["colors"]
        fonts = config["fonts"]

        # 顶部深蓝色背景区域（通过表格模拟）
        header_table = doc.add_table(rows=1, cols=1)
        header_cell = header_table.cell(0, 0)
        # 设置单元格背景色
        self._set_cell_background(header_cell, colors["primary"])

        # 姓名
        name_para = header_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.font.size = Pt(fonts["name_size"])
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(255, 255, 255)

        # 联系方式
        contact_parts = []
        if personal_info.get("phone"):
            contact_parts.append(f"📱 {personal_info['phone']}")
        if personal_info.get("email"):
            contact_parts.append(f"✉️ {personal_info['email']}")
        if personal_info.get("address"):
            contact_parts.append(f"📍 {personal_info['address']}")

        if contact_parts:
            contact_para = header_cell.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run("  |  ".join(contact_parts))
            contact_run.font.size = Pt(fonts["body_size"])
            contact_run.font.color.rgb = RGBColor(200, 200, 200)

        doc.add_paragraph()  # 空行

        # 个人简介
        if summary:
            self._add_section_header_business(doc, "个人简介", colors)
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(fonts["body_size"])

        # 工作经历
        if experience:
            self._add_section_header_business(doc, "工作经历", colors)
            for exp in experience:
                self._add_experience_item(doc, exp, fonts, colors)

        # 教育背景
        if education:
            self._add_section_header_business(doc, "教育背景", colors)
            for edu in education:
                self._add_education_item(doc, edu, fonts, colors)

        # 专业技能
        if skills:
            self._add_section_header_business(doc, "专业技能", colors)
            skills_para = doc.add_paragraph()
            for i, skill in enumerate(skills):
                skills_para.add_run(f"✓ {skill}  ")
                if (i + 1) % 3 == 0 and i < len(skills) - 1:
                    skills_para = doc.add_paragraph()
            for run in skills_para.runs:
                run.font.size = Pt(fonts["body_size"])

        # 项目经历
        if projects:
            self._add_section_header_business(doc, "项目经历", colors)
            for proj in projects:
                self._add_project_item(doc, proj, fonts, colors)

        # 获奖荣誉
        if awards:
            self._add_section_header_business(doc, "获奖荣誉", colors)
            for award in awards:
                p = doc.add_paragraph(f"🏆 {award}")
                for run in p.runs:
                    run.font.size = Pt(fonts["body_size"])

    def _build_creative_resume(
        self,
        doc: Document,
        config: dict,
        personal_info: dict,
        summary: str,
        education: list,
        experience: list,
        skills: list,
        projects: list,
        awards: list,
    ) -> None:
        """构建创意风格简历（彩色侧边栏效果）。"""
        colors = config["colors"]
        fonts = config["fonts"]

        # 创建两列布局表格
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # 设置列宽
        left_col = table.columns[0]
        right_col = table.columns[1]
        for cell in left_col.cells:
            cell.width = Cm(5.5)
        for cell in right_col.cells:
            cell.width = Cm(12)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # 左侧彩色区域
        self._set_cell_background(left_cell, colors["primary"])

        # 左侧内容：姓名
        name_para = left_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.font.size = Pt(fonts["name_size"])
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(255, 255, 255)

        # 左侧：联系方式
        if personal_info.get("phone"):
            p = left_cell.add_paragraph(f"📱 {personal_info['phone']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

        if personal_info.get("email"):
            p = left_cell.add_paragraph(f"✉️ {personal_info['email']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

        if personal_info.get("address"):
            p = left_cell.add_paragraph(f"📍 {personal_info['address']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

        # 左侧：技能
        if skills:
            left_cell.add_paragraph()
            skill_header = left_cell.add_paragraph("技能")
            skill_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in skill_header.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(12)

            for skill in skills:
                p = left_cell.add_paragraph(f"• {skill}")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)

        # 右侧内容
        # 个人简介
        if summary:
            header = right_cell.paragraphs[0]
            header_run = header.add_run("个人简介")
            header_run.font.bold = True
            header_run.font.size = Pt(fonts["section_size"])
            header_run.font.color.rgb = colors["secondary"]

            p = right_cell.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(fonts["body_size"])
            right_cell.add_paragraph()

        # 工作经历
        if experience:
            header = right_cell.add_paragraph()
            header_run = header.add_run("工作经历")
            header_run.font.bold = True
            header_run.font.size = Pt(fonts["section_size"])
            header_run.font.color.rgb = colors["secondary"]

            for exp in experience:
                title_para = right_cell.add_paragraph()
                title_run = title_para.add_run(f"{exp.get('title', '')} @ {exp.get('company', '')}")
                title_run.font.bold = True
                title_run.font.size = Pt(11)
                title_run.font.color.rgb = colors["primary"]

                date_para = right_cell.add_paragraph()
                date_run = date_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                date_run.font.size = Pt(10)
                date_run.font.color.rgb = colors["accent"]

                if exp.get("description"):
                    desc_para = right_cell.add_paragraph(exp["description"])
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])

        # 教育经历
        if education:
            header = right_cell.add_paragraph()
            header_run = header.add_run("教育经历")
            header_run.font.bold = True
            header_run.font.size = Pt(fonts["section_size"])
            header_run.font.color.rgb = colors["secondary"]

            for edu in education:
                title_para = right_cell.add_paragraph()
                title_run = title_para.add_run(f"{edu.get('school', '')} - {edu.get('degree', '')} {edu.get('major', '')}")
                title_run.font.bold = True
                title_run.font.size = Pt(11)

                if edu.get("start_date") or edu.get("end_date"):
                    date_para = right_cell.add_paragraph()
                    date_run = date_para.add_run(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = colors["accent"]

        # 项目经历
        if projects:
            header = right_cell.add_paragraph()
            header_run = header.add_run("项目经历")
            header_run.font.bold = True
            header_run.font.size = Pt(fonts["section_size"])
            header_run.font.color.rgb = colors["secondary"]

            for proj in projects:
                title_para = right_cell.add_paragraph()
                title_run = title_para.add_run(proj.get("name", ""))
                title_run.font.bold = True
                title_run.font.color.rgb = colors["primary"]

                if proj.get("description"):
                    desc_para = right_cell.add_paragraph(proj["description"])
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])

    def _build_academic_resume(
        self,
        doc: Document,
        config: dict,
        personal_info: dict,
        summary: str,
        education: list,
        experience: list,
        skills: list,
        projects: list,
        awards: list,
    ) -> None:
        """构建学术风格简历。"""
        colors = config["colors"]
        fonts = config["fonts"]

        # 姓名（居中，宋体）
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.font.size = Pt(fonts["name_size"])
        name_run.font.bold = True
        name_run.font.name = fonts["name"]

        # 联系方式（居中）
        contact_parts = []
        if personal_info.get("phone"):
            contact_parts.append(f"电话: {personal_info['phone']}")
        if personal_info.get("email"):
            contact_parts.append(f"邮箱: {personal_info['email']}")
        if personal_info.get("address"):
            contact_parts.append(f"地址: {personal_info['address']}")

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(fonts["body_size"])
            contact_run.font.name = fonts["name"]

        self._add_horizontal_line(doc)

        # 教育背景（学术简历首位）
        if education:
            self._add_section_header_academic(doc, "教育背景", colors, fonts)
            for edu in education:
                edu_para = doc.add_paragraph()
                # 学校名称加粗
                school_run = edu_para.add_run(f"{edu.get('school', '')}")
                school_run.font.bold = True
                school_run.font.size = Pt(fonts["body_size"])
                school_run.font.name = fonts["name"]

                # 学位和专业
                degree_run = edu_para.add_run(f"，{edu.get('degree', '')} {edu.get('major', '')}")
                degree_run.font.size = Pt(fonts["body_size"])
                degree_run.font.name = fonts["name"]

                # 日期和GPA
                details = []
                if edu.get("start_date") or edu.get("end_date"):
                    details.append(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
                if edu.get("gpa"):
                    details.append(f"GPA: {edu['gpa']}")

                if details:
                    detail_para = doc.add_paragraph("    " + " | ".join(details))
                    for run in detail_para.runs:
                        run.font.size = Pt(10)
                        run.font.name = fonts["name"]
                        run.font.color.rgb = colors["secondary"]

        # 研究/工作经历
        if experience:
            self._add_section_header_academic(doc, "研究/工作经历", colors, fonts)
            for exp in experience:
                exp_para = doc.add_paragraph()
                title_run = exp_para.add_run(f"{exp.get('title', '')}")
                title_run.font.bold = True
                title_run.font.size = Pt(fonts["body_size"])
                title_run.font.name = fonts["name"]

                company_run = exp_para.add_run(f"，{exp.get('company', '')}")
                company_run.font.size = Pt(fonts["body_size"])
                company_run.font.name = fonts["name"]

                if exp.get("start_date") or exp.get("end_date"):
                    date_para = doc.add_paragraph(f"    {exp.get('start_date', '')} - {exp.get('end_date', '')}")
                    for run in date_para.runs:
                        run.font.size = Pt(10)
                        run.font.name = fonts["name"]
                        run.font.color.rgb = colors["secondary"]

                if exp.get("description"):
                    desc_para = doc.add_paragraph(f"    {exp['description']}")
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])
                        run.font.name = fonts["name"]

        # 项目/研究成果
        if projects:
            self._add_section_header_academic(doc, "研究项目", colors, fonts)
            for proj in projects:
                proj_para = doc.add_paragraph()
                name_run = proj_para.add_run(f"• {proj.get('name', '')}")
                name_run.font.bold = True
                name_run.font.size = Pt(fonts["body_size"])
                name_run.font.name = fonts["name"]

                if proj.get("description"):
                    desc_para = doc.add_paragraph(f"    {proj['description']}")
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])
                        run.font.name = fonts["name"]

        # 获奖荣誉
        if awards:
            self._add_section_header_academic(doc, "获奖与荣誉", colors, fonts)
            for award in awards:
                p = doc.add_paragraph(f"• {award}")
                for run in p.runs:
                    run.font.size = Pt(fonts["body_size"])
                    run.font.name = fonts["name"]

        # 专业技能
        if skills:
            self._add_section_header_academic(doc, "专业技能", colors, fonts)
            skills_para = doc.add_paragraph("、".join(skills))
            for run in skills_para.runs:
                run.font.size = Pt(fonts["body_size"])
                run.font.name = fonts["name"]

    def _build_technical_resume(
        self,
        doc: Document,
        config: dict,
        personal_info: dict,
        summary: str,
        education: list,
        experience: list,
        skills: list,
        projects: list,
        awards: list,
    ) -> None:
        """构建技术风格简历（带技能进度条）。"""
        colors = config["colors"]
        fonts = config["fonts"]

        # 顶部区域
        header_table = doc.add_table(rows=1, cols=1)
        header_cell = header_table.cell(0, 0)
        self._set_cell_background(header_cell, colors["secondary"])

        name_para = header_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", ""))
        name_run.font.size = Pt(fonts["name_size"])
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(255, 255, 255)
        name_run.font.name = fonts["name"]

        # 联系方式
        contact_parts = []
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("address"):
            contact_parts.append(personal_info["address"])

        if contact_parts:
            contact_para = header_cell.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(200, 200, 200)
            contact_run.font.name = fonts["name"]

        doc.add_paragraph()

        # 个人简介
        if summary:
            self._add_section_header_tech(doc, "// 个人简介", colors)
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(fonts["body_size"])

        # 技术栈（带进度条效果）
        if skills:
            self._add_section_header_tech(doc, "// 技术栈", colors)
            # 使用表格模拟进度条
            for skill in skills:
                skill_table = doc.add_table(rows=1, cols=2)
                skill_table.autofit = False

                name_cell = skill_table.cell(0, 0)
                name_cell.width = Cm(4)
                name_para = name_cell.paragraphs[0]
                name_run = name_para.add_run(skill)
                name_run.font.size = Pt(10)
                name_run.font.name = fonts["name"]

                # 进度条单元格
                bar_cell = skill_table.cell(0, 1)
                bar_cell.width = Cm(10)
                self._set_cell_background(bar_cell, colors["primary"])
                bar_para = bar_cell.paragraphs[0]
                bar_run = bar_para.add_run("█" * 10)
                bar_run.font.color.rgb = RGBColor(255, 255, 255)
                bar_run.font.size = Pt(8)

        # 工作经历
        if experience:
            self._add_section_header_tech(doc, "// 工作经历", colors)
            for exp in experience:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"[{exp.get('company', '')}] {exp.get('title', '')}")
                title_run.font.bold = True
                title_run.font.size = Pt(11)
                title_run.font.color.rgb = colors["primary"]
                title_run.font.name = fonts["name"]

                if exp.get("start_date") or exp.get("end_date"):
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(f"📅 {exp.get('start_date', '')} - {exp.get('end_date', '')}")
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = colors["accent"]

                if exp.get("description"):
                    desc_para = doc.add_paragraph(exp["description"])
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])

        # 项目经历
        if projects:
            self._add_section_header_tech(doc, "// 项目经历", colors)
            for proj in projects:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"📦 {proj.get('name', '')}")
                title_run.font.bold = True
                title_run.font.size = Pt(11)
                title_run.font.color.rgb = colors["primary"]

                if proj.get("role"):
                    role_para = doc.add_paragraph()
                    role_run = role_para.add_run(f"角色: {proj['role']}")
                    role_run.font.size = Pt(10)
                    role_run.font.color.rgb = colors["secondary"]

                if proj.get("technologies"):
                    tech_para = doc.add_paragraph()
                    tech_run = tech_para.add_run(f"技术栈: {proj['technologies']}")
                    tech_run.font.size = Pt(10)
                    tech_run.font.color.rgb = colors["accent"]

                if proj.get("description"):
                    desc_para = doc.add_paragraph(proj["description"])
                    for run in desc_para.runs:
                        run.font.size = Pt(fonts["body_size"])

        # 教育背景
        if education:
            self._add_section_header_tech(doc, "// 教育背景", colors)
            for edu in education:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"🎓 {edu.get('school', '')} | {edu.get('degree', '')} {edu.get('major', '')}")
                edu_run.font.size = Pt(11)
                edu_run.font.name = fonts["name"]

                if edu.get("start_date") or edu.get("end_date"):
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(f"   {edu.get('start_date', '')} - {edu.get('end_date', '')}")
                    date_run.font.size = Pt(10)
                    date_run.font.color.rgb = colors["accent"]

        # 获奖荣誉
        if awards:
            self._add_section_header_tech(doc, "// 获奖荣誉", colors)
            for award in awards:
                p = doc.add_paragraph(f"🏆 {award}")
                for run in p.runs:
                    run.font.size = Pt(fonts["body_size"])

    # ========== 辅助方法 ==========

    def _add_horizontal_line(self, doc: Document) -> None:
        """添加水平分隔线。"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("─" * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)

    def _add_section_header(
        self, doc: Document, title: str, color: RGBColor, size: int
    ) -> None:
        """添加通用章节标题。"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = color

    def _add_section_header_business(
        self, doc: Document, title: str, colors: dict
    ) -> None:
        """添加商务风格章节标题（带下划线）。"""
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

    def _add_section_header_academic(
        self, doc: Document, title: str, colors: dict, fonts: dict
    ) -> None:
        """添加学术风格章节标题。"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(fonts["section_size"])
        run.font.color.rgb = colors["accent"]
        run.font.name = fonts["name"]
        # 下划线效果
        underline_p = doc.add_paragraph()
        underline_run = underline_p.add_run("─" * 30)
        underline_run.font.color.rgb = colors["accent"]

    def _add_section_header_tech(
        self, doc: Document, title: str, colors: dict
    ) -> None:
        """添加技术风格章节标题（代码注释风格）。"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]
        run.font.name = "Consolas"

    def _add_experience_item(
        self, doc: Document, exp: dict, fonts: dict, colors: dict
    ) -> None:
        """添加工作经历条目。"""
        # 职位 @ 公司
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{exp.get('title', '')} @ {exp.get('company', '')}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        # 日期
        if exp.get("start_date") or exp.get("end_date"):
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = colors["secondary"]

        # 描述
        if exp.get("description"):
            desc_para = doc.add_paragraph(exp["description"])
            for run in desc_para.runs:
                run.font.size = Pt(fonts["body_size"])

    def _add_education_item(
        self, doc: Document, edu: dict, fonts: dict, colors: dict
    ) -> None:
        """添加教育经历条目。"""
        # 学校 - 学位 专业
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{edu.get('school', '')} - {edu.get('degree', '')} {edu.get('major', '')}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        # 日期和GPA
        details = []
        if edu.get("start_date") or edu.get("end_date"):
            details.append(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
        if edu.get("gpa"):
            details.append(f"GPA: {edu['gpa']}")

        if details:
            detail_para = doc.add_paragraph()
            detail_run = detail_para.add_run(" | ".join(details))
            detail_run.font.size = Pt(10)
            detail_run.font.color.rgb = colors["secondary"]

    def _add_project_item(
        self, doc: Document, proj: dict, fonts: dict, colors: dict
    ) -> None:
        """添加项目经历条目。"""
        # 项目名称
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(proj.get("name", ""))
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        # 角色
        if proj.get("role"):
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"担任角色: {proj['role']}")
            role_run.font.size = Pt(10)
            role_run.font.color.rgb = colors["secondary"]

        # 技术栈
        if proj.get("technologies"):
            tech_para = doc.add_paragraph()
            tech_run = tech_para.add_run(f"技术栈: {proj['technologies']}")
            tech_run.font.size = Pt(10)
            tech_run.font.color.rgb = colors["accent"]

        # 描述
        if proj.get("description"):
            desc_para = doc.add_paragraph(proj["description"])
            for run in desc_para.runs:
                run.font.size = Pt(fonts["body_size"])

    def _set_cell_background(self, cell, color: RGBColor) -> None:
        """设置单元格背景色。"""
        if color is None:
            return
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color.hex()}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _export_resume(self, params: dict[str, Any]) -> ToolResult:
        """导出简历为其他格式。"""
        input_file = params.get("input_file", "")
        output_format = params.get("format", "pdf").lower()

        if not input_file:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="必须指定输入文件路径 (input_file)",
            )

        input_path = Path(input_file)
        if not input_path.exists():
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"输入文件不存在: {input_file}",
            )

        if not input_path.suffix.lower() == ".docx":
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"输入文件必须是DOCX格式，当前: {input_path.suffix}",
            )

        if output_format not in ["pdf", "docx", "html"]:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的导出格式: {output_format}，支持: pdf/docx/html",
            )

        # 如果是docx格式，直接复制
        if output_format == "docx":
            output_path = self.output_dir / f"{input_path.stem}_exported.docx"
            shutil.copy(input_path, output_path)
            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=f"✅ 简历已导出\n📁 文件: {output_path}",
                data={"file_path": str(output_path), "format": "docx"},
            )

        try:
            if output_format == "pdf":
                return self._export_to_pdf(input_path)
            elif output_format == "html":
                return self._export_to_html(input_path)
        except Exception as e:
            logger.error("导出失败: %s", e, exc_info=True)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"导出失败: {e}",
            )

    def _export_to_pdf(self, input_path: Path) -> ToolResult:
        """导出为PDF（使用Pandoc + XeLaTeX）。"""
        # 检查pandoc是否可用
        if not shutil.which("pandoc"):
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="Pandoc未安装，无法导出PDF。请安装Pandoc: https://pandoc.org/installing.html",
            )

        output_path = self.output_dir / f"{input_path.stem}.pdf"

        try:
            # 使用pandoc转换，指定xelatex引擎以支持中文
            result = subprocess.run(
                [
                    "pandoc",
                    str(input_path),
                    "-o",
                    str(output_path),
                    "--pdf-engine=xelatex",
                    "-V",
                    "mainfont=SimSun",
                    "-V",
                    "CJKmainfont=SimSun",
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                # 尝试不带中文字体选项
                result = subprocess.run(
                    ["pandoc", str(input_path), "-o", str(output_path)],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )

                if result.returncode != 0:
                    return ToolResult(
                        status=ToolResultStatus.ERROR,
                        error=f"PDF导出失败: {result.stderr}",
                    )

            file_size = output_path.stat().st_size if output_path.exists() else 0

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=f"✅ 简历已导出为PDF\n📁 文件: {output_path.name}\n📊 大小: {file_size} 字节",
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "format": "pdf",
                },
            )
        except subprocess.TimeoutExpired:
            return ToolResult(
                status=ToolResultStatus.TIMEOUT,
                error="PDF导出超时",
            )
        except Exception as e:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"PDF导出失败: {e}",
            )

    def _export_to_html(self, input_path: Path) -> ToolResult:
        """导出为HTML（使用内联CSS确保样式）。"""
        output_path = self.output_dir / f"{input_path.stem}.html"

        try:
            # 读取DOCX内容
            doc = Document(str(input_path))

            # 构建HTML
            html_parts = [
                "<!DOCTYPE html>",
                "<html>",
                "<head>",
                '<meta charset="utf-8">',
                f"<title>{input_path.stem}</title>",
                "<style>",
                self._get_html_styles(),
                "</style>",
                "</head>",
                "<body>",
                '<div class="resume-container">',
            ]

            # 转换文档内容
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else "2"
                    html_parts.append(f"<h{level}>{html.escape(para.text)}</h{level}>")
                elif para.text.strip():
                    # 处理格式
                    text = html.escape(para.text)
                    for run in para.runs:
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                    html_parts.append(f"<p>{text}</p>")

            # 处理表格
            for table in doc.tables:
                html_parts.append("<table>")
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        html_parts.append(f"<td>{html.escape(cell.text)}</td>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html_parts.extend(["</div>", "</body>", "</html>"])

            # 写入文件
            output_path.write_text("\n".join(html_parts), encoding="utf-8")
            file_size = output_path.stat().st_size

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=f"✅ 简历已导出为HTML\n📁 文件: {output_path.name}\n📊 大小: {file_size} 字节",
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "format": "html",
                },
            )
        except Exception as e:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"HTML导出失败: {e}",
            )

    def _get_html_styles(self) -> str:
        """获取HTML内联样式。"""
        return """
            * { margin: 0; padding: 0; box-sizing: border-box; }
            body {
                font-family: 'Microsoft YaHei', 'Segoe UI', Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                background: #f5f5f5;
                padding: 20px;
            }
            .resume-container {
                max-width: 800px;
                margin: 0 auto;
                background: white;
                padding: 40px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            h1 {
                font-size: 28px;
                color: #2c3e50;
                text-align: center;
                margin-bottom: 10px;
                border-bottom: 2px solid #3498db;
                padding-bottom: 10px;
            }
            h2 {
                font-size: 18px;
                color: #2980b9;
                margin: 20px 0 10px;
                padding-bottom: 5px;
                border-bottom: 1px solid #eee;
            }
            h3 {
                font-size: 14px;
                color: #34495e;
                margin: 15px 0 5px;
            }
            p {
                margin: 8px 0;
                text-align: justify;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 10px 0;
            }
            td, th {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            th {
                background: #f8f9fa;
            }
            strong { color: #2c3e50; }
            em { color: #7f8c8d; }
        """


# 用于测试
if __name__ == "__main__":
    import asyncio

    async def test():
        tool = ResumeBuilderTool()

        # 测试列出模板
        result = await tool.execute("list_templates", {})
        print("=== 模板列表 ===")
        print(result.output)
        print()

        # 测试生成简历
        result = await tool.execute(
            "generate_resume",
            {
                "template": "technical",
                "personal_info": {
                    "name": "张三",
                    "phone": "13800138000",
                    "email": "zhangsan@example.com",
                    "address": "北京市海淀区",
                },
                "summary": "5年Python开发经验，熟悉Web开发、数据分析和机器学习。",
                "education": [
                    {
                        "school": "清华大学",
                        "degree": "硕士",
                        "major": "计算机科学",
                        "start_date": "2015.09",
                        "end_date": "2018.06",
                        "gpa": "3.8/4.0",
                    }
                ],
                "experience": [
                    {
                        "company": "字节跳动",
                        "title": "高级软件工程师",
                        "start_date": "2020.07",
                        "end_date": "至今",
                        "description": "负责推荐系统后端开发，优化算法性能提升30%。",
                    }
                ],
                "skills": ["Python", "Go", "Kubernetes", "TensorFlow", "React"],
                "projects": [
                    {
                        "name": "智能推荐引擎",
                        "role": "技术负责人",
                        "technologies": "Python, TensorFlow, Redis",
                        "description": "设计并实现了基于深度学习的推荐系统，日均处理1亿次请求。",
                    }
                ],
                "awards": ["2022年公司技术创新奖", "ACM竞赛省级一等奖"],
            },
        )
        print("=== 生成结果 ===")
        print(result.output)
        print("Data:", result.data)

    asyncio.run(test())
