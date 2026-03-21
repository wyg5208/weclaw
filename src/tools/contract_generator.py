"""合同生成工具 — 智能合同文档生成工具。

支持生成租赁合同、劳动合同、买卖合同、服务合同等常见合同类型。
"""

from __future__ import annotations

import logging
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from jinja2 import Template

from src.tools.base import ActionDef, BaseTool, ToolResult, ToolResultStatus

logger = logging.getLogger(__name__)


# 合同模板定义
CONTRACT_TEMPLATES = {
    "rental": {
        "name": "租赁合同",
        "description": "房屋/物品租赁合同，适用于出租方与承租方之间的租赁关系",
        "required_fields": ["租赁物", "租期", "租金", "押金", "支付方式"],
        "default_clauses": {
            "租赁物条款": "甲方将位于{{ terms.property_location | default('___') }}的房屋/物品出租给乙方使用。",
            "租期条款": "租赁期限自{{ terms.start_date | default('___年___月___日') }}起至{{ terms.end_date | default('___年___月___日') }}止，共计{{ terms.duration | default('___') }}个月。",
            "租金条款": "租金为每月人民币{{ terms.rent | default('___') }}元整（大写：{{ terms.rent_cn | default('___') }}），乙方应于每月{{ terms.pay_day | default('___') }}日前支付当月租金。",
            "押金条款": "乙方应于签订本合同时向甲方支付押金人民币{{ terms.deposit | default('___') }}元整，合同期满且无违约情况下，甲方应于{{ terms.return_days | default('15') }}个工作日内退还押金。",
            "支付方式条款": "租金及押金通过{{ terms.payment_method | default('银行转账') }}方式支付至甲方指定账户。",
            "维护条款": "乙方应合理使用租赁物，承担日常维护责任。因乙方使用不当造成的损坏，由乙方负责修复或赔偿。",
            "解约条款": "任何一方提前解除合同，应提前{{ terms.notice_days | default('30') }}天书面通知对方，否则应支付违约金。",
        },
    },
    "labor": {
        "name": "劳动合同",
        "description": "用人单位与劳动者之间的劳动合同，明确双方权利义务",
        "required_fields": ["岗位", "工资", "试用期", "工作时间", "社保"],
        "default_clauses": {
            "岗位条款": "甲方聘用乙方担任{{ terms.position | default('___') }}岗位，工作地点为{{ terms.work_location | default('___') }}。",
            "工资条款": "乙方月工资为人民币{{ terms.salary | default('___') }}元整，甲方于每月{{ terms.pay_day | default('___') }}日前发放上月工资。",
            "试用期条款": "试用期为{{ terms.probation_months | default('___') }}个月，试用期工资为正式工资的{{ terms.probation_rate | default('80') }}%。",
            "工作时间条款": "乙方实行{{ terms.work_system | default('标准工时') }}制，每日工作{{ terms.work_hours | default('8') }}小时，每周工作{{ terms.work_days | default('5') }}天。",
            "社保条款": "甲方按照国家规定为乙方缴纳社会保险，包括养老保险、医疗保险、失业保险、工伤保险、生育保险。",
            "保密条款": "乙方应对甲方的商业秘密和技术秘密保密，离职后{{ terms.confidential_years | default('2') }}年内不得向第三方披露。",
            "解除条款": "任何一方解除合同应提前{{ terms.notice_days | default('30') }}天书面通知对方，双方协商一致可随时解除。",
        },
    },
    "sales": {
        "name": "买卖合同",
        "description": "商品买卖合同，适用于买方与卖方之间的交易",
        "required_fields": ["商品名称", "数量", "单价", "总价", "交货方式"],
        "default_clauses": {
            "商品条款": "甲方向乙方出售{{ terms.product_name | default('___') }}，规格型号为{{ terms.product_spec | default('___') }}。",
            "数量条款": "商品数量为{{ terms.quantity | default('___') }}{{ terms.unit | default('件') }}，总价为人民币{{ terms.total_price | default('___') }}元整。",
            "价格条款": "商品单价为人民币{{ terms.unit_price | default('___') }}元/{{ terms.unit | default('件') }}，价格包含{{ terms.price_includes | default('运费') }}。",
            "交货条款": "甲方应于{{ terms.delivery_date | default('___年___月___日') }}前将商品交付至{{ terms.delivery_location | default('___') }}。",
            "付款条款": "乙方应于{{ terms.payment_time | default('收货后___天内') }}通过{{ terms.payment_method | default('银行转账') }}方式支付货款。",
            "验收条款": "乙方应在收货后{{ terms.inspection_days | default('3') }}天内完成验收，逾期未提出异议视为验收合格。",
            "质保条款": "甲方对商品提供{{ terms.warranty_months | default('12') }}个月质保服务，质保期内非人为损坏免费维修或更换。",
        },
    },
    "service": {
        "name": "服务合同",
        "description": "服务委托合同，适用于委托方与服务方之间的服务关系",
        "required_fields": ["服务内容", "服务费用", "服务期限", "验收标准"],
        "default_clauses": {
            "服务内容条款": "乙方为甲方提供{{ terms.service_content | default('___') }}服务，服务范围包括{{ terms.service_scope | default('___') }}。",
            "服务期限条款": "服务期限自{{ terms.start_date | default('___年___月___日') }}起至{{ terms.end_date | default('___年___月___日') }}止。",
            "费用条款": "服务费用为人民币{{ terms.service_fee | default('___') }}元整，{{ terms.payment_terms | default('分期支付') }}。",
            "验收条款": "甲方应在乙方提交服务成果后{{ terms.acceptance_days | default('5') }}个工作日内完成验收，验收标准为{{ terms.acceptance_criteria | default('___') }}。",
            "保密条款": "乙方应对服务过程中获知的甲方信息保密，未经甲方书面同意不得向第三方披露。",
            "违约条款": "任何一方违约应向对方支付合同总额{{ terms.penalty_rate | default('10') }}%的违约金。",
            "知识产权条款": "服务过程中产生的知识产权归{{ terms.ip_owner | default('甲方') }}所有。",
        },
    },
}


class ContractGeneratorTool(BaseTool):
    """合同生成工具。

    支持生成多种类型的合同文档，包括租赁合同、劳动合同、买卖合同、服务合同。
    支持自定义条款和多格式导出。
    """

    name = "contract_generator"
    emoji = "📄"
    title = "合同生成器"
    description = "智能合同文档生成工具"
    timeout = 120

    def __init__(self, output_dir: str = None) -> None:
        """初始化合同生成工具。

        Args:
            output_dir: 输出目录，默认为项目的 generated/日期/ 目录
        """
        super().__init__()
        self.output_dir = (
            Path(output_dir)
            if output_dir
            else Path(__file__).parent.parent.parent / "generated" / datetime.now().strftime("%Y-%m-%d")
        )
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 自定义条款存储（运行时）
        self._custom_clauses: dict[str, dict[str, str]] = {}

    def get_actions(self) -> list[ActionDef]:
        return [
            ActionDef(
                name="generate_contract",
                description=(
                    "根据合同类型和参数生成合同文档。支持租赁合同(rental)、劳动合同(labor)、"
                    "买卖合同(sales)、服务合同(service)四种类型。生成完整的DOCX格式合同文件。"
                ),
                parameters={
                    "contract_type": {
                        "type": "string",
                        "description": "合同类型: rental(租赁)、labor(劳动)、sales(买卖)、service(服务)",
                        "enum": ["rental", "labor", "sales", "service"],
                    },
                    "party_a": {
                        "type": "object",
                        "description": "甲方信息",
                        "properties": {
                            "name": {"type": "string", "description": "甲方名称/姓名"},
                            "address": {"type": "string", "description": "甲方地址"},
                            "id_number": {"type": "string", "description": "甲方身份证号/统一社会信用代码"},
                        },
                    },
                    "party_b": {
                        "type": "object",
                        "description": "乙方信息",
                        "properties": {
                            "name": {"type": "string", "description": "乙方名称/姓名"},
                            "address": {"type": "string", "description": "乙方地址"},
                            "id_number": {"type": "string", "description": "乙方身份证号/统一社会信用代码"},
                        },
                    },
                    "terms": {
                        "type": "object",
                        "description": "合同关键条款，如金额、期限、支付方式等，字段因合同类型而异",
                    },
                    "title": {
                        "type": "string",
                        "description": "合同标题（可选），不提供则使用默认标题",
                    },
                },
                required_params=["contract_type", "party_a", "party_b", "terms"],
            ),
            ActionDef(
                name="list_templates",
                description="列出所有可用的合同模板，包括模板类型、名称、描述和所需字段说明。",
                parameters={},
                required_params=[],
            ),
            ActionDef(
                name="customize_clause",
                description="自定义合同条款内容。修改指定合同模板的特定条款，支持Jinja2模板语法。",
                parameters={
                    "contract_type": {
                        "type": "string",
                        "description": "合同类型: rental、labor、sales、service",
                        "enum": ["rental", "labor", "sales", "service"],
                    },
                    "clause_name": {
                        "type": "string",
                        "description": "条款名称，如'租金条款'、'工资条款'等",
                    },
                    "content": {
                        "type": "string",
                        "description": "条款内容，支持Jinja2模板语法，可使用{{ terms.xxx }}引用参数",
                    },
                },
                required_params=["contract_type", "clause_name", "content"],
            ),
            ActionDef(
                name="export_contract",
                description="将已生成的合同DOCX文件导出为其他格式。PDF导出需要安装Pandoc和XeLaTeX。",
                parameters={
                    "input_file": {
                        "type": "string",
                        "description": "输入的DOCX文件路径",
                    },
                    "format": {
                        "type": "string",
                        "description": "导出格式: pdf、docx、txt",
                        "enum": ["pdf", "docx", "txt"],
                    },
                },
                required_params=["input_file", "format"],
            ),
        ]

    async def execute(self, action: str, params: dict[str, Any]) -> ToolResult:
        """执行指定动作。"""
        if action == "generate_contract":
            return self._generate_contract(params)
        elif action == "list_templates":
            return self._list_templates()
        elif action == "customize_clause":
            return self._customize_clause(params)
        elif action == "export_contract":
            return self._export_contract(params)
        else:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的动作: {action}",
            )

    def _generate_contract(self, params: dict[str, Any]) -> ToolResult:
        """生成合同文档。"""
        contract_type = params.get("contract_type", "").strip()
        party_a = params.get("party_a", {})
        party_b = params.get("party_b", {})
        terms = params.get("terms", {})
        title = params.get("title", "").strip()

        # 验证合同类型
        if contract_type not in CONTRACT_TEMPLATES:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的合同类型: {contract_type}，支持: {', '.join(CONTRACT_TEMPLATES.keys())}",
            )

        template = CONTRACT_TEMPLATES[contract_type]
        contract_name = template["name"]
        
        # 设置合同标题
        if not title:
            title = contract_name

        try:
            # 创建文档
            doc = Document()
            
            # 设置页面边距
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3)
                section.right_margin = Cm(3)

            # 添加合同标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_run.font.name = "SimHei"
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加合同编号和日期
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contract_no = datetime.now().strftime("合同编号：%Y%m%d%H%M%S")
            info_run = info_para.add_run(contract_no)
            info_run.font.size = Pt(10)
            
            doc.add_paragraph()  # 空行

            # 添加双方信息
            self._add_party_info(doc, "甲方", party_a)
            self._add_party_info(doc, "乙方", party_b)
            
            doc.add_paragraph()  # 空行

            # 添加合同前言
            preamble = doc.add_paragraph()
            preamble_text = (
                f"甲乙双方本着平等自愿、诚实信用的原则，经友好协商，"
                f"就{self._get_contract_subject(contract_type)}事宜达成如下协议："
            )
            preamble.add_run(preamble_text).font.size = Pt(12)
            
            doc.add_paragraph()  # 空行

            # 获取条款（合并默认条款和自定义条款）
            clauses = dict(template["default_clauses"])
            if contract_type in self._custom_clauses:
                clauses.update(self._custom_clauses[contract_type])

            # 添加合同条款
            article_num = 1
            for clause_name, clause_template in clauses.items():
                # 渲染条款内容
                try:
                    jinja_template = Template(clause_template)
                    clause_content = jinja_template.render(terms=terms)
                except Exception as e:
                    logger.warning(f"渲染条款 {clause_name} 失败: {e}")
                    clause_content = clause_template

                # 添加条款标题
                clause_title = doc.add_paragraph()
                clause_title_run = clause_title.add_run(f"第{self._num_to_chinese(article_num)}条 {clause_name}")
                clause_title_run.bold = True
                clause_title_run.font.size = Pt(12)
                
                # 添加条款内容
                clause_para = doc.add_paragraph()
                clause_para.add_run(clause_content).font.size = Pt(12)
                clause_para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                
                article_num += 1

            # 添加争议解决条款
            dispute_title = doc.add_paragraph()
            dispute_title_run = dispute_title.add_run(f"第{self._num_to_chinese(article_num)}条 争议解决")
            dispute_title_run.bold = True
            dispute_title_run.font.size = Pt(12)
            
            dispute_para = doc.add_paragraph()
            dispute_text = "本合同在履行过程中发生争议，由双方协商解决；协商不成的，可向合同签订地人民法院提起诉讼。"
            dispute_para.add_run(dispute_text).font.size = Pt(12)
            dispute_para.paragraph_format.first_line_indent = Cm(0.74)
            article_num += 1

            # 添加生效条款
            effect_title = doc.add_paragraph()
            effect_title_run = effect_title.add_run(f"第{self._num_to_chinese(article_num)}条 合同生效")
            effect_title_run.bold = True
            effect_title_run.font.size = Pt(12)
            
            effect_para = doc.add_paragraph()
            effect_text = "本合同一式两份，甲乙双方各执一份，自双方签字（盖章）之日起生效。"
            effect_para.add_run(effect_text).font.size = Pt(12)
            effect_para.paragraph_format.first_line_indent = Cm(0.74)

            doc.add_paragraph()  # 空行
            doc.add_paragraph()  # 空行

            # 添加签字区域
            self._add_signature_area(doc)

            # 生成文件名并保存
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in "_ ")[:20]
            filename = f"contract_{contract_type}_{safe_title}_{timestamp}.docx"
            output_path = self.output_dir / filename
            
            doc.save(str(output_path))
            
            file_size = output_path.stat().st_size
            
            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=(
                    f"✅ 合同已生成\n"
                    f"📋 类型: {contract_name}\n"
                    f"📁 文件: {output_path.name}\n"
                    f"📊 大小: {file_size} 字节\n"
                    f"📍 路径: {output_path}"
                ),
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "contract_type": contract_type,
                    "title": title,
                },
            )
        except Exception as e:
            logger.error(f"生成合同失败: {e}", exc_info=True)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"生成合同失败: {e}",
            )

    def _add_party_info(self, doc: Document, party_label: str, party_info: dict) -> None:
        """添加当事人信息。"""
        name = party_info.get("name", "_______________")
        address = party_info.get("address", "_______________")
        id_number = party_info.get("id_number", "_______________")

        para = doc.add_paragraph()
        para.add_run(f"{party_label}（{'出租方' if party_label == '甲方' else '承租方'}）：").bold = True
        para.add_run(name).font.size = Pt(12)
        
        addr_para = doc.add_paragraph()
        addr_para.add_run("地址：").font.size = Pt(12)
        addr_para.add_run(address).font.size = Pt(12)
        
        id_para = doc.add_paragraph()
        id_para.add_run("身份证号/统一社会信用代码：").font.size = Pt(12)
        id_para.add_run(id_number).font.size = Pt(12)

    def _add_signature_area(self, doc: Document) -> None:
        """添加签字区域。"""
        # 创建两列签字区域
        table = doc.add_table(rows=4, cols=2)
        table.autofit = True
        
        # 甲方签字区
        table.cell(0, 0).text = "甲方（签字/盖章）："
        table.cell(1, 0).text = ""
        table.cell(2, 0).text = ""
        table.cell(3, 0).text = f"日期：____年____月____日"
        
        # 乙方签字区
        table.cell(0, 1).text = "乙方（签字/盖章）："
        table.cell(1, 1).text = ""
        table.cell(2, 1).text = ""
        table.cell(3, 1).text = f"日期：____年____月____日"

    def _get_contract_subject(self, contract_type: str) -> str:
        """获取合同主题描述。"""
        subjects = {
            "rental": "房屋/物品租赁",
            "labor": "建立劳动关系",
            "sales": "商品买卖",
            "service": "服务委托",
        }
        return subjects.get(contract_type, "相关")

    def _num_to_chinese(self, num: int) -> str:
        """数字转中文。"""
        chinese_nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                        "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]
        if 0 <= num <= 20:
            return chinese_nums[num]
        return str(num)

    def _list_templates(self) -> ToolResult:
        """列出所有可用的合同模板。"""
        output_lines = ["📋 可用合同模板列表：", ""]
        
        for type_key, template in CONTRACT_TEMPLATES.items():
            output_lines.append(f"【{template['name']}】({type_key})")
            output_lines.append(f"  描述: {template['description']}")
            output_lines.append(f"  必需字段: {', '.join(template['required_fields'])}")
            output_lines.append(f"  条款数量: {len(template['default_clauses'])}条")
            output_lines.append("")
        
        templates_data = []
        for type_key, template in CONTRACT_TEMPLATES.items():
            templates_data.append({
                "type": type_key,
                "name": template["name"],
                "description": template["description"],
                "required_fields": template["required_fields"],
                "clauses": list(template["default_clauses"].keys()),
            })
        
        return ToolResult(
            status=ToolResultStatus.SUCCESS,
            output="\n".join(output_lines),
            data={"templates": templates_data},
        )

    def _customize_clause(self, params: dict[str, Any]) -> ToolResult:
        """自定义合同条款。"""
        contract_type = params.get("contract_type", "").strip()
        clause_name = params.get("clause_name", "").strip()
        content = params.get("content", "").strip()

        if contract_type not in CONTRACT_TEMPLATES:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的合同类型: {contract_type}",
            )

        if not clause_name:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="条款名称不能为空",
            )

        if not content:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="条款内容不能为空",
            )

        # 验证Jinja2模板语法
        try:
            Template(content)
        except Exception as e:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"条款模板语法错误: {e}",
            )

        # 存储自定义条款
        if contract_type not in self._custom_clauses:
            self._custom_clauses[contract_type] = {}
        self._custom_clauses[contract_type][clause_name] = content

        return ToolResult(
            status=ToolResultStatus.SUCCESS,
            output=(
                f"✅ 条款已自定义\n"
                f"📋 合同类型: {CONTRACT_TEMPLATES[contract_type]['name']}\n"
                f"📝 条款名称: {clause_name}\n"
                f"💡 提示: 下次生成该类型合同时将使用自定义条款"
            ),
            data={
                "contract_type": contract_type,
                "clause_name": clause_name,
                "content": content,
            },
        )

    def _export_contract(self, params: dict[str, Any]) -> ToolResult:
        """导出合同为不同格式。"""
        input_file = params.get("input_file", "").strip()
        output_format = params.get("format", "").strip().lower()

        if not input_file:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="输入文件路径不能为空",
            )

        input_path = Path(input_file)
        if not input_path.exists():
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"输入文件不存在: {input_file}",
            )

        if not input_path.suffix.lower() == ".docx":
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="输入文件必须是DOCX格式",
            )

        if output_format not in ["pdf", "docx", "txt"]:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的导出格式: {output_format}，支持: pdf, docx, txt",
            )

        try:
            output_path = input_path.with_suffix(f".{output_format}")

            if output_format == "docx":
                # 复制文件
                import shutil
                shutil.copy2(input_path, output_path)
                engine = "copy"

            elif output_format == "txt":
                # 提取文本
                doc = Document(str(input_path))
                text_content = []
                for para in doc.paragraphs:
                    text_content.append(para.text)
                output_path.write_text("\n".join(text_content), encoding="utf-8")
                engine = "python-docx"

            elif output_format == "pdf":
                # 使用 Pandoc + XeLaTeX 导出 PDF
                if not shutil.which("pandoc"):
                    return ToolResult(
                        status=ToolResultStatus.ERROR,
                        error="PDF导出需要安装Pandoc。请访问 https://pandoc.org/installing.html 安装。",
                    )

                # 创建临时文件
                with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as tmp_md:
                    # 从DOCX提取文本并写入临时Markdown文件
                    doc = Document(str(input_path))
                    for para in doc.paragraphs:
                        if para.style and para.style.name and para.style.name.startswith("Heading"):
                            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                            tmp_md.write(f"{'#' * level} {para.text}\n\n")
                        else:
                            tmp_md.write(f"{para.text}\n\n")
                    tmp_md_path = tmp_md.name

                try:
                    # 使用 Pandoc 转换为 PDF
                    result = subprocess.run(
                        [
                            "pandoc",
                            tmp_md_path,
                            "-o", str(output_path),
                            "--pdf-engine=xelatex",
                            "-V", "CJKmainfont=SimSun",
                            "-V", "geometry:margin=2.5cm",
                        ],
                        capture_output=True,
                        text=True,
                        timeout=60,
                    )
                    if result.returncode != 0:
                        # 尝试不使用 XeLaTeX
                        result = subprocess.run(
                            ["pandoc", tmp_md_path, "-o", str(output_path)],
                            capture_output=True,
                            text=True,
                            timeout=60,
                        )
                        if result.returncode != 0:
                            raise Exception(f"Pandoc转换失败: {result.stderr}")
                    engine = "pandoc+xelatex"
                finally:
                    # 清理临时文件
                    try:
                        Path(tmp_md_path).unlink()
                    except Exception:
                        pass

            file_size = output_path.stat().st_size if output_path.exists() else 0

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=(
                    f"✅ 合同已导出\n"
                    f"📁 文件: {output_path.name}\n"
                    f"📊 大小: {file_size} 字节\n"
                    f"⚙️ 引擎: {engine}\n"
                    f"📍 路径: {output_path}"
                ),
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "format": output_format,
                    "engine_used": engine,
                },
            )

        except subprocess.TimeoutExpired:
            return ToolResult(
                status=ToolResultStatus.TIMEOUT,
                error="PDF导出超时，请检查XeLaTeX是否正确安装",
            )
        except Exception as e:
            logger.error(f"导出合同失败: {e}", exc_info=True)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"导出合同失败: {e}",
            )


# 用于测试
if __name__ == "__main__":
    import asyncio

    async def test():
        tool = ContractGeneratorTool()
        
        # 测试列出模板
        print("=== 测试列出模板 ===")
        result = await tool.execute("list_templates", {})
        print(result.output)
        print()

        # 测试生成租赁合同
        print("=== 测试生成租赁合同 ===")
        result = await tool.execute("generate_contract", {
            "contract_type": "rental",
            "party_a": {
                "name": "张三",
                "address": "北京市朝阳区XX街道XX号",
                "id_number": "110101199001011234",
            },
            "party_b": {
                "name": "李四",
                "address": "北京市海淀区XX街道XX号",
                "id_number": "110102199201021234",
            },
            "terms": {
                "property_location": "北京市朝阳区XX小区XX号楼XX单元XXX室",
                "start_date": "2026年4月1日",
                "end_date": "2027年3月31日",
                "duration": "12",
                "rent": "5000",
                "rent_cn": "伍仟元整",
                "pay_day": "5",
                "deposit": "10000",
                "return_days": "15",
                "payment_method": "银行转账",
                "notice_days": "30",
            },
            "title": "房屋租赁合同",
        })
        print(result.output)
        print()

        # 测试自定义条款
        print("=== 测试自定义条款 ===")
        result = await tool.execute("customize_clause", {
            "contract_type": "rental",
            "clause_name": "特别约定条款",
            "content": "甲乙双方特别约定：{{ terms.special_terms | default('无') }}。",
        })
        print(result.output)

    asyncio.run(test())
