"""文档生成工具 — 将 Markdown 内容转换为 Word 或 HTML 文档。
"""

from __future__ import annotations

import html
import logging
import subprocess
import tempfile
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from src.tools.base import ActionDef, BaseTool, ToolResult, ToolResultStatus

logger = logging.getLogger(__name__)


class HTMLToDocxParser(HTMLParser):
    """将 HTML 转换为 python-docx 文档的解析器。"""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_runs = []  # 当前段落中的 runs

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)
        if tag in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[tag]
            self.doc.add_heading("", level=level)
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "p":
            self.doc.add_paragraph()
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "br":
            # HTML <br> 在段落中表现为换行，这里简单处理为新段落
            self.doc.add_paragraph()
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "b" or tag == "strong":
            # 标记当前需要加粗
            self.current_bold = True
        elif tag == "i" or tag == "em":
            # 标记当前需要斜体
            self.current_italic = True
        elif tag == "code":
            self.current_mono = True
        elif tag == "ul":
            self.in_list = "ul"
        elif tag == "ol":
            self.in_list = "ol"
            self.list_counter = 0
        elif tag == "li":
            if self.in_list == "ul":
                self.doc.add_paragraph("", style="List Bullet")
            elif self.in_list == "ol":
                self.list_counter += 1
                p = self.doc.add_paragraph("", style="List Number")
                # 修改编号为实际数字（List Number 样式通常从1开始）
                pass
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "blockquote":
            self.in_blockquote = True
            self.doc.add_paragraph()
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "pre":
            self.in_pre = True
            self.doc.add_paragraph()
            self.current_runs = self.doc.paragraphs[-1].runs
        elif tag == "img":
            src = attrs_dict.get("src", "")
            alt = attrs_dict.get("alt", "")
            if src:
                self._add_image(src, alt)

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "p", "li", "blockquote", "pre"):
            self.current_runs = []
        elif tag == "b" or tag == "strong":
            self.current_bold = False
        elif tag == "i" or tag == "em":
            self.current_italic = False
        elif tag == "code":
            self.current_mono = False
        elif tag == "ul" or tag == "ol":
            self.in_list = None
        elif tag == "blockquote":
            self.in_blockquote = False
            # 设置引用样式
            if self.doc.paragraphs:
                p = self.doc.paragraphs[-1]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
        elif tag == "pre":
            self.in_pre = False
            # 设置代码块样式
            if self.doc.paragraphs:
                p = self.doc.paragraphs[-1]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)

    def handle_data(self, data: str) -> None:
        if not data.strip():
            return
            
        # 创建新的 run
        if not self.doc.paragraphs:
            self.doc.add_paragraph()
            self.current_runs = self.doc.paragraphs[-1].runs
            
        if not self.current_runs:
            p = self.doc.paragraphs[-1] if self.doc.paragraphs else self.doc.add_paragraph()
            run = p.add_run(data)
            self.current_runs = [run]
        else:
            run = self.current_runs[-1] if self.current_runs else self.doc.paragraphs[-1].add_run(data)
            if len(self.current_runs) == 0 or run.text:
                run = self.doc.paragraphs[-1].add_run(data)
                self.current_runs.append(run)
            else:
                run.text += data

        # 应用格式
        if hasattr(self, "current_bold") and self.current_bold:
            run.bold = True
        if hasattr(self, "current_italic") and self.current_italic:
            run.italic = True
        if hasattr(self, "current_mono") and self.current_mono:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        if hasattr(self, "in_blockquote") and self.in_blockquote:
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        if hasattr(self, "in_pre") and self.in_pre:
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    def _add_image(self, src: str, alt: str = "") -> None:
        """添加图片到文档。"""
        try:
            # 处理相对路径
            img_path = Path(src)
            if not img_path.is_absolute():
                img_path = Path.cwd() / img_path
                
            if img_path.exists():
                # 添加图片说明
                if alt:
                    caption = self.doc.add_paragraph(alt)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                        
                # 添加图片
                img_paragraph = self.doc.add_paragraph()
                run = img_paragraph.add_run()
                run.add_picture(str(img_path), width=Inches(4.5))
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.info("已添加图片: %s", img_path)
            else:
                # 图片不存在，添加占位文本
                p = self.doc.add_paragraph(f"[图片: {alt or src}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.warning("添加图片失败: %s - %s", src, e)
            p = self.doc.add_paragraph(f"[图片加载失败: {alt or src}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


class DocGeneratorTool(BaseTool):
    """文档生成工具。

    支持将 Markdown 内容转换为 Word (.docx) 或 HTML 文档。
    采用双引擎策略：优先使用 Pandoc，失败后降级到 python-docx。
    """

    name = "doc_generator"
    emoji = "📄"
    title = "文档生成"
    description = "生成 Word 文档或 HTML 文件，支持 Markdown 格式输入"
    timeout = 60

    def __init__(self, output_dir: str = "") -> None:
        """初始化文档生成工具。

        Args:
            output_dir: 输出目录，默认为项目的 generated/日期/ 目录
        """
        super().__init__()
        # 直接使用带日期的子目录，避免重复复制
        self.output_dir = Path(output_dir) if output_dir else Path(__file__).parent.parent.parent / "generated" / datetime.now().strftime("%Y-%m-%d")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 初始化状态标志
        self.current_bold = False
        self.current_italic = False
        self.current_mono = False
        self.in_list = None
        self.list_counter = 0
        self.in_blockquote = False
        self.in_pre = False

    def get_actions(self) -> list[ActionDef]:
        return [
            ActionDef(
                name="generate_document",
                description=(
                    "生成文档。支持 Markdown 格式输入，可输出为 Word (.docx) 或 HTML 格式。"
                    "自动支持标题、列表、代码块、引用、表格、图片等语法。"
                ),
                parameters={
                    "content": {
                        "type": "string",
                        "description": "文档内容，支持 Markdown 格式",
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题，默认为'AI生成文档'",
                    },
                    "format_type": {
                        "type": "string",
                        "description": "输出格式: docx(默认) 或 html",
                        "enum": ["docx", "html"],
                    },
                    "filename": {
                        "type": "string",
                        "description": "自定义文件名（建议添加主题），如: doc_诗歌一首_20260215_135033。不提供则自动生成 doc_年月日_时分秒",
                    },
                },
                required_params=["content"],
            ),
            ActionDef(
                name="generate_from_template",
                description=(
                    "从专业模板生成文档。支持 academic_paper（学术论文）、business_report（商业报告）、"
                    "meeting_minutes（会议纪要）、contract（合同）、resume（简历）等模板。"
                ),
                parameters={
                    "template": {
                        "type": "string",
                        "description": "模板名称",
                        "enum": ["academic_paper", "business_report", "meeting_minutes", "contract", "resume"],
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题/主题",
                    },
                    "author": {
                        "type": "string",
                        "description": "作者/姓名",
                    },
                    "organization": {
                        "type": "string",
                        "description": "组织/公司名称",
                    },
                    "date": {
                        "type": "string",
                        "description": "日期，默认当前日期",
                    },
                    "color_scheme": {
                        "type": "string",
                        "description": "配色方案: business/academic/tech/nature",
                        "enum": ["business", "academic", "tech", "nature"],
                    },
                    "filename": {
                        "type": "string",
                        "description": "输出文件名（不含扩展名）",
                    },
                },
                required_params=["template", "title"],
            ),
            ActionDef(
                name="list_templates",
                description="列出所有可用的文档模板",
                parameters={},
                required_params=[],
            ),
            ActionDef(
                name="validate_document",
                description="验证 Word 文档的结构完整性（检查表格、图片、书签等）",
                parameters={
                    "file_path": {
                        "type": "string",
                        "description": "要验证的 DOCX 文件路径",
                    },
                },
                required_params=["file_path"],
            ),
        ]

    async def execute(self, action: str, params: dict[str, Any]) -> ToolResult:
        if action == "generate_document":
            return self._generate_document(params)
        elif action == "generate_from_template":
            return self._generate_from_template(params)
        elif action == "list_templates":
            return self._list_templates(params)
        elif action == "validate_document":
            return self._validate_document(params)
        else:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的动作: {action}",
            )

    def _generate_document(self, params: dict[str, Any]) -> ToolResult:
        """生成文档的核心逻辑。"""
        content = params.get("content", "").strip()
        if not content:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="文档内容不能为空",
            )

        title = params.get("title", "AI生成文档").strip()
        format_type = params.get("format_type", "docx").lower()
        filename = params.get("filename", "").strip()

        # 生成文件名
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"doc_{timestamp}"
        output_path = self.output_dir / f"{filename}.{format_type}"

        try:
            if format_type == "html":
                engine = self._generate_html(content, title, output_path)
            elif format_type == "docx":
                engine = self._generate_docx(content, title, output_path)
            else:
                return ToolResult(
                    status=ToolResultStatus.ERROR,
                    error=f"不支持的格式: {format_type}，支持 docx/html",
                )

            file_size = output_path.stat().st_size if output_path.exists() else 0
            
            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=f"✅ 文档已生成\n📁 文件: {output_path.name}\n📊 大小: {file_size} 字节\n⚙️  引擎: {engine}",
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "format_type": format_type,
                    "engine_used": engine,
                    "title": title,
                },
            )
        except Exception as e:
            logger.error("文档生成失败: %s", e, exc_info=True)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"文档生成失败: {e}",
            )

    def _generate_html(self, content: str, title: str, output_path: Path) -> str:
        """生成 HTML 文档。"""
        # 使用 markdown 库转换，启用表格和代码块扩展
        html_body = markdown.markdown(content, extensions=["tables", "fenced_code"])
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{html.escape(title)}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 800px;
            margin: 40px auto;
            line-height: 1.6;
            color: #333;
            padding: 0 20px;
        }}
        h1, h2, h3 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
        }}
        h1 {{
            font-size: 2em;
            border-bottom: 1px solid #eee;
            padding-bottom: 0.3em;
        }}
        h2 {{
            font-size: 1.5em;
            border-bottom: 1px solid #eee;
            padding-bottom: 0.3em;
        }}
        h3 {{
            font-size: 1.25em;
        }}
        p {{
            margin: 0 0 16px 0;
        }}
        ul, ol {{
            padding-left: 2em;
            margin-bottom: 16px;
        }}
        li {{
            margin-bottom: 4px;
        }}
        blockquote {{
            margin: 0 0 16px 0;
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
        }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            padding: 16px;
            overflow-x: auto;
            margin-bottom: 16px;
        }}
        code {{
            font-family: 'SFMono-Regular', Consolas, 'Liberation Mono', Menlo, monospace;
            font-size: 85%;
            background-color: rgba(27,31,35,0.05);
            border-radius: 3px;
            padding: 0.2em 0.4em;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 16px;
        }}
        th, td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        img {{
            max-width: 100%;
            box-sizing: content-box;
        }}
    </style>
</head>
<body>
    <h1>{html.escape(title)}</h1>
    {html_body}
</body>
</html>"""
        
        output_path.write_text(html_content, encoding="utf-8")
        logger.info("HTML 文档已生成: %s", output_path)
        return "markdown+html"

    def _generate_docx(self, content: str, title: str, output_path: Path) -> str:
        """生成 Word 文档（使用 Pandoc）。"""
        try:
            # 检查 pandoc 是否可用
            import shutil
            if not shutil.which('pandoc'):
                raise Exception("Pandoc 未安装")
            
            # 创建临时 Markdown 文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
                tmp_md.write(content)
                tmp_md_path = tmp_md.name
            
            try:
                # 使用 subprocess 调用 pandoc
                import subprocess
                result = subprocess.run(
                    ['pandoc', tmp_md_path, '-o', str(output_path)],
                    capture_output=True,
                    text=True
                )
                if result.returncode != 0:
                    raise Exception(f"Pandoc 转换失败: {result.stderr}")
                
                logger.info("Pandoc 文档生成成功: %s", output_path)
                return "pandoc"
            finally:
                # 清理临时文件
                try:
                    import os
                    os.unlink(tmp_md_path)
                except:
                    pass
        except Exception as e:
            logger.warning("Pandoc 转换失败，使用 python-docx 降级: %s", e)
            # 降级到 python-docx
            return self._fallback_docx_from_markdown(content, title, output_path)

    def _fallback_docx_from_markdown(self, content: str, title: str, output_path: Path) -> str:
        """使用 python-docx 直接从 Markdown 生成文档（不经过 Pandoc）。"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import re
        
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, level=1)
        
        # 创建 Code 样式（如果不存在）
        try:
            code_style = doc.styles.get_or_create_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)
        except Exception:
            pass  # 样式创建失败不影响主流程
        
        # 解析 Markdown 内容
        lines = content.split('\n')
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []
        
        for line in lines:
            line = line.rstrip()
            
            # 处理代码块
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    # 代码块结束，添加到文档
                    code = '\n'.join(code_content)
                    p = doc.add_paragraph(code)
                    try:
                        p.style = 'Code'
                    except KeyError:
                        # 样式不存在时，直接设置等宽字体
                        p.runs[0].font.name = 'Courier New'
                        p.runs[0].font.size = Pt(10)
                    in_code_block = False
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # 处理表格
            if line.startswith('|'):
                # 跳过表格分隔行 |---|---|
                if re.match(r'^\|[\s\-:|]+\|$', line):
                    continue
                # 解析表格行
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                table_rows.append(cells)
                in_table = True
                continue
            elif in_table:
                # 表格结束，生成表格
                if table_rows:
                    self._add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            
            # 处理图片：![alt](path)
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                
                # 检查图片路径是否存在
                from pathlib import Path as PathLib
                img_file = PathLib(img_path)
                if not img_file.is_absolute():
                    img_file = PathLib.cwd() / img_path
                
                if img_file.exists():
                    try:
                        # 添加图片
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(str(img_file), width=Inches(4.5))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 添加图片说明
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in caption.runs:
                                r.font.size = Pt(10)
                                r.font.color.rgb = RGBColor(100, 100, 100)
                    except Exception as e:
                        logger.warning("添加图片失败: %s - %s", img_path, e)
                        doc.add_paragraph(f"[图片加载失败: {alt_text or img_path}]")
                else:
                    # 图片文件不存在
                    doc.add_paragraph(f"[图片: {alt_text or img_path}]")
                continue
            
            # 处理标题
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                # 普通段落
                doc.add_paragraph(line)
        
        # 处理文件末尾的表格
        if in_table and table_rows:
            self._add_table_to_doc(doc, table_rows)
        
        # 保存文档
        doc.save(str(output_path))
        logger.info("python-docx 文档生成成功: %s", output_path)
        return "python-docx"
    
    def _add_table_to_doc(self, doc: Document, table_rows: list) -> None:
        """将 Markdown 表格数据添加到 Word 文档。"""
        if not table_rows:
            return
        
        # 创建表格
        num_cols = len(table_rows[0]) if table_rows else 0
        if num_cols == 0:
            return
        
        # 如果第一行看起来像表头（没有特殊符号），则用它作为表头
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    # 表头行加粗
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _fallback_docx(self, html_content: str, title: str, output_path: Path) -> str:
        """使用 python-docx 生成文档（降级方案）。"""
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, level=1)
        
        # 解析 HTML 并构建 docx
        parser = HTMLToDocxParser(doc)
        parser.feed(html_content)
        
        # 保存文档
        doc.save(output_path)
        logger.info("python-docx 生成成功: %s", output_path)
        return "python-docx"

    def _generate_from_template(self, params: dict[str, Any]) -> ToolResult:
        """从模板生成文档。"""
        template_name = params.get("template", "").strip()
        if not template_name:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="模板名称不能为空",
            )

        title = params.get("title", "").strip()
        if not title:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="文档标题不能为空",
            )

        author = params.get("author", "").strip()
        organization = params.get("organization", "").strip()
        date = params.get("date", "").strip() or datetime.now().strftime("%Y年%m月%d日")
        color_scheme = params.get("color_scheme", "business")
        filename = params.get("filename", "").strip()

        # 验证模板名称
        valid_templates = ["academic_paper", "business_report", "meeting_minutes", "contract", "resume"]
        if template_name not in valid_templates:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"无效的模板: {template_name}，可用模板: {', '.join(valid_templates)}",
            )

        # 验证配色方案
        valid_schemes = ["business", "academic", "tech", "nature"]
        if color_scheme not in valid_schemes:
            color_scheme = "business"

        try:
            # 导入模板模块
            import sys
            from pathlib import Path as PathLib
            sys.path.insert(0, str(PathLib(__file__).parent.parent.parent))
            from resources.docx_templates.templates import (
                TemplateOptions,
                create_from_template,
            )

            options = TemplateOptions(
                title=title,
                author=author,
                date=date,
                organization=organization,
                color_scheme=color_scheme,
            )

            doc = create_from_template(template_name, options)
            if not doc:
                return ToolResult(
                    status=ToolResultStatus.ERROR,
                    error=f"模板创建失败: {template_name}",
                )

            # 生成文件名
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{template_name}_{timestamp}"

            output_path = self.output_dir / f"{filename}.docx"
            doc.save(str(output_path))

            file_size = output_path.stat().st_size

            logger.info("模板文档生成成功: %s", output_path)

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=f"✅ 模板文档生成完成\n📁 文件: {output_path.name}\n📊 大小: {file_size} 字节\n📋 模板: {template_name}\n🎨 配色: {color_scheme}",
                data={
                    "file_path": str(output_path),
                    "file_name": output_path.name,
                    "file_size": file_size,
                    "template": template_name,
                    "color_scheme": color_scheme,
                    "title": title,
                },
            )

        except Exception as e:
            logger.error("模板文档生成失败: %s", e)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"模板文档生成失败: {e}",
            )

    def _list_templates(self, params: dict[str, Any]) -> ToolResult:
        """列出所有可用模板。"""
        try:
            import sys
            from pathlib import Path as PathLib
            sys.path.insert(0, str(PathLib(__file__).parent.parent.parent))
            from resources.docx_templates.templates import list_templates

            templates = list_templates()

            output = "📋 可用文档模板:\n\n"
            for tmpl in templates:
                output += f"• {tmpl['title']} ({tmpl['name']})\n"
                output += f"  {tmpl['description']}\n\n"

            return ToolResult(
                status=ToolResultStatus.SUCCESS,
                output=output,
                data={"templates": templates},
            )

        except Exception as e:
            logger.error("列出模板失败: %s", e)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"列出模板失败: {e}",
            )

    def _validate_document(self, params: dict[str, Any]) -> ToolResult:
        """验证 Word 文档的结构完整性。"""
        file_path = params.get("file_path", "").strip()
        if not file_path:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error="文件路径不能为空",
            )

        docx_path = Path(file_path)
        if not docx_path.exists():
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"文件不存在: {file_path}",
            )

        if docx_path.suffix.lower() not in (".docx", ".docm"):
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"不支持的文件格式: {docx_path.suffix}",
            )

        try:
            # 使用 DOCX 验证器
            from src.tools.docx_validator import run_validation

            report = run_validation(str(docx_path))

            if report.passed:
                output = f"✅ 文档验证通过\n📁 文件: {docx_path.name}"
                if report.warnings:
                    output += f"\n⚠️  警告 ({len(report.warnings)}):\n"
                    for warn in report.warnings[:5]:
                        output += f"  • {warn}\n"
            else:
                output = f"❌ 文档验证失败\n📁 文件: {docx_path.name}\n\n"
                if report.errors:
                    output += "错误:\n"
                    for err in report.errors:
                        output += f"  • {err}\n"
                if report.warnings:
                    output += "\n警告:\n"
                    for warn in report.warnings[:5]:
                        output += f"  • {warn}\n"

            return ToolResult(
                status=ToolResultStatus.SUCCESS if report.passed else ToolResultStatus.ERROR,
                output=output,
                data={
                    "passed": report.passed,
                    "errors": report.errors,
                    "warnings": report.warnings,
                },
            )

        except ImportError:
            # 降级：使用基本验证
            return self._basic_validate_docx(docx_path)
        except Exception as e:
            logger.error("文档验证失败: %s", e)
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"文档验证失败: {e}",
            )

    def _basic_validate_docx(self, docx_path: Path) -> ToolResult:
        """基本 DOCX 验证（不依赖 docx_validator）。"""
        try:
            doc = Document(str(docx_path))

            # 统计信息
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            section_count = len(doc.sections)

            # 检查内容
            has_content = para_count > 0

            output = f"📊 文档分析:\n"
            output += f"  • 段落数: {para_count}\n"
            output += f"  • 表格数: {table_count}\n"
            output += f"  • 节数: {section_count}\n"

            if has_content:
                output += f"\n✅ 文档结构正常"
                status = ToolResultStatus.SUCCESS
            else:
                output += f"\n⚠️ 文档可能为空"
                status = ToolResultStatus.WARNING

            return ToolResult(
                status=status,
                output=output,
                data={
                    "passed": has_content,
                    "paragraph_count": para_count,
                    "table_count": table_count,
                    "section_count": section_count,
                },
            )

        except Exception as e:
            return ToolResult(
                status=ToolResultStatus.ERROR,
                error=f"文档验证失败: {e}",
            )


# 用于测试
if __name__ == "__main__":
    import asyncio
    
    async def test():
        tool = DocGeneratorTool()
        result = await tool.execute("generate_document", {
            "content": """# 测试文档

这是一个**加粗**和*斜体*的测试。

## 代码块
```python
print("Hello World")
```

## 列表
- 项目1
- 项目2
  - 子项目2.1

1. 第一项
2. 第二项

## 引用
> 这是一段引用文本

## 表格
| 姓名 | 年龄 |
|------|------|
| 张三 | 25   |
| 李四 | 30   |
""",
            "title": "测试文档",
            "format_type": "docx"
        })
        print(result.output)
        print("Data:", result.data)
    
    asyncio.run(test())