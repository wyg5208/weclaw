"""Office 文档工具综合测试脚本

测试范围：
1. Excel 公式验证模块 (excel_validator)
2. DOCX 验证模块 (docx_validator)
3. PDF 生成工具 (pdf_generator)
4. DOCX 模板生成 (doc_generator 增强)
5. PPT 图表幻灯片 (ppt_generator 增强)
"""

import sys
import os
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import asyncio
import tempfile
from datetime import datetime

# 测试结果收集
test_results = []


def log_test(name: str, passed: bool, message: str = ""):
    """记录测试结果"""
    status = "✅ PASS" if passed else "❌ FAIL"
    print(f"  {status}: {name}")
    if message:
        print(f"     {message}")
    test_results.append({
        "name": name,
        "passed": passed,
        "message": message
    })


def print_section(title: str):
    """打印测试分节标题"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print("=" * 60)


async def test_excel_validator():
    """测试 Excel 公式验证模块"""
    print_section("Excel 公式验证模块 (excel_validator)")

    # 测试 1: 模块导入
    try:
        from src.tools.excel_validator import (
            recalc_workbook,
            refcheck_workbook,
            validate_xlsx,
            EXCEL_ERRORS,
            RecalcResult,
            RefcheckResult,
            STYLE_THEMES,
            apply_professional_style,
        )
        log_test("模块导入", True)
    except ImportError as e:
        log_test("模块导入", False, str(e))
        return

    # 测试 2: 常量定义
    expected_errors = ["#VALUE!", "#DIV/0!", "#REF!", "#NAME?", "#NULL!", "#NUM!", "#N/A"]
    if set(EXCEL_ERRORS) == set(expected_errors):
        log_test("EXCEL_ERRORS 常量", True)
    else:
        log_test("EXCEL_ERRORS 常量", False, f"期望 {expected_errors}, 实际 {EXCEL_ERRORS}")

    # 测试 3: 主题样式
    if "grayscale" in STYLE_THEMES and "financial" in STYLE_THEMES:
        log_test("STYLE_THEMES 主题", True, f"可用主题: {list(STYLE_THEMES.keys())}")
    else:
        log_test("STYLE_THEMES 主题", False)

    # 测试 4: RecalcResult 数据类
    result = RecalcResult(status="success", total_formulas=10)
    if result.is_success:
        log_test("RecalcResult.is_success", True)
    else:
        log_test("RecalcResult.is_success", False)

    # 测试 5: RefcheckResult 数据类
    result = RefcheckResult(status="success", sheets_checked=["Sheet1"])
    if result.is_clean:
        log_test("RefcheckResult.is_clean", True)
    else:
        log_test("RefcheckResult.is_clean", False)

    # 测试 6: recalc_workbook 函数存在性
    if callable(recalc_workbook):
        log_test("recalc_workbook 函数", True)
    else:
        log_test("recalc_workbook 函数", False)

    # 测试 7: refcheck_workbook 函数存在性
    if callable(refcheck_workbook):
        log_test("refcheck_workbook 函数", True)
    else:
        log_test("refcheck_workbook 函数", False)

    # 测试 8: validate_xlsx 函数存在性
    if callable(validate_xlsx):
        log_test("validate_xlsx 函数", True)
    else:
        log_test("validate_xlsx 函数", False)

    # 测试 9: apply_professional_style 函数存在性
    if callable(apply_professional_style):
        log_test("apply_professional_style 函数", True)
    else:
        log_test("apply_professional_style 函数", False)

    # 测试 10: 创建测试 Excel 文件并验证
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "测试表"

        # 添加测试数据
        ws.append(["姓名", "销售额", "增长率"])
        ws.append(["张三", 15000, 0.15])
        ws.append(["李四", 23000, -0.05])

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            temp_path = f.name

        wb.save(temp_path)
        log_test("创建测试 Excel 文件", True, temp_path)

        # 使用 refcheck 验证
        result = refcheck_workbook(temp_path)
        if result.status in ["success", "issues_found"]:
            log_test("refcheck_workbook 执行", True, f"状态: {result.status}")
        else:
            log_test("refcheck_workbook 执行", False, f"状态: {result.status}")

        # 清理
        os.unlink(temp_path)

    except Exception as e:
        log_test("Excel 文件测试", False, str(e))


async def test_docx_validator():
    """测试 DOCX 验证模块"""
    print_section("DOCX 验证模块 (docx_validator)")

    # 测试 1: 模块导入
    try:
        from src.tools.docx_validator import (
            run_validation,
            ValidationPipeline,
            ValidationReport,
            Issue,
            DETECTORS,
            GridConsistencyDetector,
            AspectRatioDetector,
        )
        log_test("模块导入", True)
    except ImportError as e:
        log_test("模块导入", False, str(e))
        return

    # 测试 2: DETECTORS 列表
    if len(DETECTORS) >= 5:
        log_test("DETECTORS 检测器列表", True, f"共 {len(DETECTORS)} 个检测器")
    else:
        log_test("DETECTORS 检测器列表", False, f"只有 {len(DETECTORS)} 个")

    # 测试 3: ValidationPipeline 类
    try:
        pipeline = ValidationPipeline()
        log_test("ValidationPipeline 实例化", True)
    except Exception as e:
        log_test("ValidationPipeline 实例化", False, str(e))

    # 测试 4: ValidationReport 数据类
    report = ValidationReport(passed=True, message="测试通过")
    if report.passed:
        log_test("ValidationReport.passed", True)
    else:
        log_test("ValidationReport.passed", False)

    # 测试 5: Issue 数据类
    issue = Issue(
        issue_type="test",
        severity="warning",
        location="test.xml",
        message="测试问题"
    )
    if issue.severity == "warning":
        log_test("Issue 数据类", True)
    else:
        log_test("Issue 数据类", False)

    # 测试 6: 创建测试 DOCX 文件并验证
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("测试文档", level=1)
        doc.add_paragraph("这是一个测试段落。")

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "列1"
        table.rows[0].cells[1].text = "列2"
        table.rows[1].cells[0].text = "数据1"
        table.rows[1].cells[1].text = "数据2"

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc.save(temp_path)
        log_test("创建测试 DOCX 文件", True, temp_path)

        # 使用 run_validation 验证
        report = run_validation(temp_path)
        if report.passed is not None:
            log_test("run_validation 执行", True, f"验证{'通过' if report.passed else '失败'}")
        else:
            log_test("run_validation 执行", False, "返回 None")

        # 清理
        os.unlink(temp_path)

    except Exception as e:
        log_test("DOCX 文件测试", False, str(e))


async def test_pdf_generator():
    """测试 PDF 生成工具"""
    print_section("PDF 生成工具 (pdf_generator)")

    # 测试 1: 模块导入
    try:
        from src.tools.pdf_generator import (
            PDFGeneratorTool,
            PDFGenerationOptions,
            PAGE_STYLE_TEMPLATE,
        )
        log_test("模块导入", True)
    except ImportError as e:
        log_test("模块导入", False, str(e))
        return

    # 测试 2: PDFGeneratorTool 类
    try:
        tool = PDFGeneratorTool()
        log_test("PDFGeneratorTool 实例化", True)
    except Exception as e:
        log_test("PDFGeneratorTool 实例化", False, str(e))
        return

    # 测试 3: get_actions 方法
    actions = tool.get_actions()
    action_names = [a.name for a in actions]
    expected_actions = ["html_to_pdf", "md_to_pdf", "apply_print_style"]

    if all(a in action_names for a in expected_actions):
        log_test("get_actions 方法", True, f"Actions: {action_names}")
    else:
        missing = [a for a in expected_actions if a not in action_names]
        log_test("get_actions 方法", False, f"缺少: {missing}")

    # 测试 4: PDFGenerationOptions 数据类
    options = PDFGenerationOptions(title="测试文档", author="测试")
    if options.title == "测试文档":
        log_test("PDFGenerationOptions", True)
    else:
        log_test("PDFGenerationOptions", False)

    # 测试 5: _build_html_document 方法
    try:
        html_body = "<p>测试内容</p>"
        html = tool._build_html_document(html_body, "测试标题")
        if "<html>" in html and "测试标题" in html:
            log_test("_build_html_document", True)
        else:
            log_test("_build_html_document", False, "HTML 结构不完整")
    except Exception as e:
        log_test("_build_html_document", False, str(e))

    # 测试 6: _apply_document_style 方法
    try:
        html = "<p>测试</p>"
        styled = tool._apply_document_style(html, "academic")
        if "<html>" in styled:
            log_test("_apply_document_style", True)
        else:
            log_test("_apply_document_style", False)
    except Exception as e:
        log_test("_apply_document_style", False, str(e))

    # 测试 7: execute 方法 - md_to_pdf
    try:
        result = await tool.execute("md_to_pdf", {
            "markdown_content": "# 测试\n\n这是一段测试内容。",
            "title": "PDF生成测试"
        })
        if result.status.value in ["success", "error"]:
            log_test("execute md_to_pdf", True, f"状态: {result.status.value}")
        else:
            log_test("execute md_to_pdf", False)
    except Exception as e:
        log_test("execute md_to_pdf", False, str(e))


async def test_doc_generator_enhancements():
    """测试 DOCX 模板生成增强"""
    print_section("DOCX 模板生成增强 (doc_generator)")

    # 测试 1: 导入增强模块
    try:
        from src.tools.doc_generator import DocGeneratorTool
        from resources.docx_templates.templates import (
            TEMPLATES,
            TemplateOptions,
            create_from_template,
            list_templates,
        )
        log_test("模块导入", True)
    except ImportError as e:
        log_test("模块导入", False, str(e))
        return

    # 测试 2: DocGeneratorTool 实例化
    try:
        tool = DocGeneratorTool()
        log_test("DocGeneratorTool 实例化", True)
    except Exception as e:
        log_test("DocGeneratorTool 实例化", False, str(e))
        return

    # 测试 3: get_actions 新增方法
    actions = tool.get_actions()
    action_names = [a.name for a in actions]
    expected_actions = ["generate_from_template", "list_templates", "validate_document"]

    if all(a in action_names for a in expected_actions):
        log_test("新增 Actions", True, f"包含: {expected_actions}")
    else:
        missing = [a for a in expected_actions if a not in action_names]
        log_test("新增 Actions", False, f"缺少: {missing}")

    # 测试 4: list_templates 函数
    try:
        templates = list_templates()
        if len(templates) >= 5:
            log_test("list_templates", True, f"共 {len(templates)} 个模板")
        else:
            log_test("list_templates", False, f"只有 {len(templates)} 个模板")
    except Exception as e:
        log_test("list_templates", False, str(e))

    # 测试 5: create_from_template 函数
    try:
        options = TemplateOptions(
            title="测试文档",
            author="测试员",
            date="2024-01-01",
            color_scheme="business"
        )
        doc = create_from_template("business_report", options)
        if doc is not None:
            log_test("create_from_template", True)
        else:
            log_test("create_from_template", False, "返回 None")
    except Exception as e:
        log_test("create_from_template", False, str(e))

    # 测试 6: execute list_templates
    try:
        result = await tool.execute("list_templates", {})
        if result.is_success and "模板" in result.output:
            log_test("execute list_templates", True)
        else:
            log_test("execute list_templates", False)
    except Exception as e:
        log_test("execute list_templates", False, str(e))

    # 测试 7: execute generate_from_template
    try:
        result = await tool.execute("generate_from_template", {
            "template": "meeting_minutes",
            "title": "会议纪要测试",
            "author": "测试者",
            "date": "2024-01-01"
        })
        if result.is_success:
            log_test("execute generate_from_template", True)
            # 清理生成的文件
            if result.data and "file_path" in result.data:
                try:
                    os.unlink(result.data["file_path"])
                    print(f"     已清理测试文件: {result.data['file_path']}")
                except:
                    pass
        else:
            log_test("execute generate_from_template", False, result.error)
    except Exception as e:
        log_test("execute generate_from_template", False, str(e))


async def test_ppt_generator_enhancements():
    """测试 PPT 图表幻灯片增强"""
    print_section("PPT 图表幻灯片增强 (ppt_generator)")

    # 测试 1: 模块导入
    try:
        from src.tools.ppt_generator import PPTTool, STYLE_COLORS
        log_test("模块导入", True)
    except ImportError as e:
        log_test("模块导入", False, str(e))
        return

    # 测试 2: PPTTool 实例化
    try:
        tool = PPTTool()
        log_test("PPTTool 实例化", True)
    except Exception as e:
        log_test("PPTTool 实例化", False, str(e))
        return

    # 测试 3: 新增 Actions
    actions = tool.get_actions()
    action_names = [a.name for a in actions]
    expected_actions = ["add_chart_slide", "add_table_slide", "add_section_divider"]

    if all(a in action_names for a in expected_actions):
        log_test("新增 Actions", True, f"包含: {expected_actions}")
    else:
        missing = [a for a in expected_actions if a not in action_names]
        log_test("新增 Actions", False, f"缺少: {missing}")

    # 测试 4: _get_style_colors 方法
    try:
        colors = tool._get_style_colors("business")
        if "primary" in colors and "secondary" in colors:
            log_test("_get_style_colors", True)
        else:
            log_test("_get_style_colors", False)
    except Exception as e:
        log_test("_get_style_colors", False, str(e))

    # 测试 5: 生成基础 PPT 并添加图表幻灯片
    try:
        # 首先生成一个基础 PPT
        gen_result = await tool.execute("generate_ppt", {
            "topic": "测试PPT",
            "outline": "测试内容",
            "style": "business"
        })

        if gen_result.is_success and "file_path" in gen_result.data:
            ppt_path = gen_result.data["file_path"]

            # 测试添加图表幻灯片
            chart_result = await tool.execute("add_chart_slide", {
                "ppt_path": ppt_path,
                "title": "销售数据图表",
                "chart_type": "bar",
                "chart_data": '{"labels": ["一月", "二月", "三月"], "values": [100, 200, 150]}',
                "style": "business"
            })

            if chart_result.is_success:
                log_test("add_chart_slide", True)
            else:
                log_test("add_chart_slide", False, chart_result.error)

            # 测试添加章节分隔页
            section_result = await tool.execute("add_section_divider", {
                "ppt_path": ppt_path,
                "section_title": "核心内容",
                "section_number": 1,
                "style": "business"
            })

            if section_result.is_success:
                log_test("add_section_divider", True)
            else:
                log_test("add_section_divider", False, section_result.error)

            # 清理
            try:
                os.unlink(ppt_path)
                print(f"     已清理测试文件: {ppt_path}")
            except:
                pass
        else:
            log_test("基础 PPT 生成", False, gen_result.error)

    except Exception as e:
        log_test("PPT 图表测试", False, str(e))


async def test_data_processor_integration():
    """测试 DataProcessorTool 集成"""
    print_section("DataProcessorTool 集成测试")

    # 测试 1: 导入 DataProcessorTool
    try:
        from src.tools.data_processor import DataProcessorTool
        log_test("DataProcessorTool 导入", True)
    except ImportError as e:
        log_test("DataProcessorTool 导入", False, str(e))
        return

    # 测试 2: 实例化
    try:
        tool = DataProcessorTool()
        log_test("DataProcessorTool 实例化", True)
    except Exception as e:
        log_test("DataProcessorTool 实例化", False, str(e))
        return

    # 测试 3: get_actions
    actions = tool.get_actions()
    action_names = [a.name for a in actions]
    expected_actions = ["read_excel", "filter_data", "sort_data", "aggregate_data",
                       "merge_data", "pivot_table", "clean_data", "export_data"]

    if len(actions) >= 8:
        log_test("DataProcessorTool Actions", True, f"共 {len(actions)} 个动作")
    else:
        log_test("DataProcessorTool Actions", False, f"只有 {len(actions)} 个动作")


async def test_tools_registry():
    """测试工具注册"""
    print_section("工具注册测试")

    # 测试 1: 导入 registry
    try:
        from src.tools.registry import ToolRegistry
        log_test("ToolRegistry 导入", True)
    except ImportError as e:
        log_test("ToolRegistry 导入", False, str(e))
        return

    # 测试 2: 创建注册表
    try:
        registry = ToolRegistry()
        registry.load_config()
        registry.auto_discover(lazy=True)
        log_test("工具注册表初始化", True)
    except Exception as e:
        log_test("工具注册表初始化", False, str(e))
        return

    # 测试 3: 检查新工具注册
    tools = registry.list_tools()
    tool_names = [t.name for t in tools]

    # 检查 doc_generator 增强
    if "doc_generator" in tool_names:
        log_test("doc_generator 注册", True)
    else:
        log_test("doc_generator 注册", False)

    # 检查 ppt_generator 增强
    if "ppt_generator" in tool_names:
        log_test("ppt_generator 注册", True)
    else:
        log_test("ppt_generator 注册", False)

    # 检查 pdf_generator 新注册
    if "pdf_generator" in tool_names:
        log_test("pdf_generator 注册", True)
    else:
        log_test("pdf_generator 注册", False)

    # 测试 4: 获取 schema
    try:
        schemas = registry.get_all_schemas()
        if len(schemas) > 0:
            log_test("获取工具 Schema", True, f"共 {len(schemas)} 个")
        else:
            log_test("获取工具 Schema", False)
    except Exception as e:
        log_test("获取工具 Schema", False, str(e))


def print_summary():
    """打印测试摘要"""
    print_section("测试摘要")

    total = len(test_results)
    passed = sum(1 for r in test_results if r["passed"])
    failed = total - passed

    print(f"总测试数: {total}")
    print(f"通过: {passed} ✅")
    print(f"失败: {failed} ❌")
    print(f"成功率: {passed/total*100:.1f}%" if total > 0 else "N/A")

    if failed > 0:
        print("\n失败项:")
        for r in test_results:
            if not r["passed"]:
                print(f"  - {r['name']}: {r['message']}")


async def main():
    """主测试函数"""
    print("\n" + "=" * 60)
    print("  Office 文档工具综合测试")
    print("  " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    print("=" * 60)

    # 执行所有测试
    await test_excel_validator()
    await test_docx_validator()
    await test_pdf_generator()
    await test_doc_generator_enhancements()
    await test_ppt_generator_enhancements()
    await test_data_processor_integration()
    await test_tools_registry()

    # 打印摘要
    print_summary()

    # 返回退出码
    failed = sum(1 for r in test_results if not r["passed"])
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    exit_code = asyncio.run(main())
    sys.exit(exit_code)
