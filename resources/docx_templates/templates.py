"""DOCX 专业模板库 - 提供多种文档模板。

模板类型：
- academic_paper: 学术论文模板（封面、目录、引用）
- business_report: 商业报告模板
- meeting_minutes: 会议纪要模板
- contract: 合同模板
- resume: 简历模板

配色方案：
- 商务稳健: 深蓝、灰、少量金色点缀
- 学术中性: 靛蓝、酒红、米白
- 科技简洁: 冷灰、青色、低饱和蓝
- 自然主题: 苔绿、土黄、暖灰
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# 配色方案
COLOR_SCHEMES = {
    "business": {
        "primary": RGBColor(0x00, 0x52, 0x8A),    # 深蓝
        "secondary": RGBColor(0x33, 0x33, 0x33),  # 灰黑
        "accent": RGBColor(0x00, 0x96, 0xD6),     # 亮蓝
        "background": RGBColor(0xFF, 0xFF, 0xFF), # 白
    },
    "academic": {
        "primary": RGBColor(0x1A, 0x23, 0x7E),    # 靛蓝
        "secondary": RGBColor(0x8B, 0x00, 0x00), # 酒红
        "accent": RGBColor(0x6A, 0x5A, 0x4A),    # 棕灰
        "background": RGBColor(0xFF, 0xFF, 0xF0), # 米白
    },
    "tech": {
        "primary": RGBColor(0x2C, 0x3E, 0x50),    # 冷灰蓝
        "secondary": RGBColor(0x00, 0x8B, 0x8B), # 青色
        "accent": RGBColor(0x41, 0x88, 0xB4),     # 低饱和蓝
        "background": RGBColor(0xF8, 0xF8, 0xF8), # 浅灰白
    },
    "nature": {
        "primary": RGBColor(0x55, 0x6B, 0x4C),    # 苔绿
        "secondary": RGBColor(0x8B, 0x77, 0x4C),  # 土黄
        "accent": RGBColor(0x6B, 0x8E, 0x7A),     # 暖灰绿
        "background": RGBColor(0xFF, 0xFF, 0xF5), # 暖白
    },
}


@dataclass
class TemplateOptions:
    """模板选项。"""
    title: str = ""
    author: str = ""
    date: str = ""
    organization: str = ""
    color_scheme: str = "business"
    page_numbers: bool = True
    header_text: str = ""


def create_document(options: TemplateOptions) -> Document:
    """创建新文档。"""
    doc = Document()
    return doc


def add_cover_page(doc: Document, options: TemplateOptions) -> Document:
    """添加封面页。"""
    colors = COLOR_SCHEMES.get(options.color_scheme, COLOR_SCHEMES["business"])
    
    # 添加顶部空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(options.title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = colors["primary"]
    
    # 添加分隔线
    doc.add_paragraph()
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run("─" * 40)
    line_run.font.color.rgb = colors["accent"]
    
    # 作者信息
    if options.author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(options.author)
        author_run.font.size = Pt(16)
        author_run.font.color.rgb = colors["secondary"]
    
    # 组织信息
    if options.organization:
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_run = org_para.add_run(options.organization)
        org_run.font.size = Pt(14)
        org_run.font.color.rgb = colors["secondary"]
    
    # 日期
    if options.date:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(options.date)
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = colors["secondary"]
    
    # 分页
    doc.add_page_break()
    
    return doc


def add_table_of_contents(doc: Document, title: str = "目录") -> Document:
    """添加目录页。"""
    # 目录标题
    toc_title = doc.add_heading(title, level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加目录提示
    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = toc_note.add_run("(请在 Word 中更新域以刷新目录)")
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_page_break()
    return doc


def add_heading_with_style(
    doc: Document,
    text: str,
    level: int = 1,
    color_scheme: str = "business"
) -> Document:
    """添加带样式的标题。"""
    colors = COLOR_SCHEMES.get(color_scheme, COLOR_SCHEMES["business"])
    
    heading = doc.add_heading(text, level=level)
    
    # 应用颜色
    for run in heading.runs:
        run.font.color.rgb = colors["primary"] if level <= 2 else colors["secondary"]
    
    return doc


def add_page_numbers(doc: Document, position: str = "bottom-center") -> Document:
    """添加页码。"""
    sections = doc.sections
    for section in sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码域
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    
    return doc


def add_header(doc: Document, text: str, color_scheme: str = "business") -> Document:
    """添加页眉。"""
    colors = COLOR_SCHEMES.get(color_scheme, COLOR_SCHEMES["business"])
    
    sections = doc.sections
    for section in sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = colors["secondary"]
    
    return doc


def create_academic_paper(options: TemplateOptions) -> Document:
    """创建学术论文模板。"""
    doc = Document()
    
    # 添加封面
    add_cover_page(doc, options)
    
    # 添加目录
    add_table_of_contents(doc)
    
    # 添加摘要提示
    abstract_heading = doc.add_heading("摘要", level=1)
    doc.add_paragraph("(请在此处添加摘要内容)")
    
    # 关键词
    keywords_para = doc.add_paragraph()
    keywords_para.add_run("关键词: ").bold = True
    keywords_para.add_run("(请在此处添加关键词)")
    
    doc.add_page_break()
    
    return doc


def create_business_report(options: TemplateOptions) -> Document:
    """创建商业报告模板。"""
    doc = Document()
    
    # 设置页眉
    add_header(doc, options.title, options.color_scheme)
    add_page_numbers(doc)
    
    # 添加封面
    add_cover_page(doc, options)
    
    # 添加目录
    add_table_of_contents(doc)
    
    # 添加执行摘要
    add_heading_with_style(doc, "执行摘要", 1, options.color_scheme)
    doc.add_paragraph("(请在此处添加执行摘要内容)")
    
    # 添加正文结构
    add_heading_with_style(doc, "1. 背景", 1, options.color_scheme)
    doc.add_paragraph("(请在此处添加背景介绍)")
    
    add_heading_with_style(doc, "2. 分析", 1, options.color_scheme)
    doc.add_paragraph("(请在此处添加分析内容)")
    
    add_heading_with_style(doc, "3. 建议", 1, options.color_scheme)
    doc.add_paragraph("(请在此处添加建议内容)")
    
    add_heading_with_style(doc, "4. 结论", 1, options.color_scheme)
    doc.add_paragraph("(请在此处添加结论)")
    
    return doc


def create_meeting_minutes(options: TemplateOptions) -> Document:
    """创建会议纪要模板。"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("会议纪要", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 会议信息表格
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("会议主题:", options.title or ""),
        ("会议时间:", options.date or ""),
        ("会议地点:", ""),
        ("主持人:", options.author or ""),
        ("记录人:", ""),
    ]
    
    for i, (label, value) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        # 加粗标签
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()
    
    # 出席人员
    heading1 = doc.add_heading("出席人员", level=2)
    doc.add_paragraph("(请列出出席人员)")
    
    # 缺席人员
    heading2 = doc.add_heading("缺席人员", level=2)
    doc.add_paragraph("(请列出缺席人员)")
    
    # 议程
    heading3 = doc.add_heading("议程", level=2)
    doc.add_paragraph("(请列出会议议程)")
    
    # 会议内容
    heading4 = doc.add_heading("会议内容", level=2)
    doc.add_paragraph("(请详细记录会议内容)")
    
    # 决议
    heading5 = doc.add_heading("决议事项", level=2)
    doc.add_paragraph("(请列出会议决议)")
    
    # 行动项
    heading6 = doc.add_heading("行动项", level=2)
    action_table = doc.add_table(rows=3, cols=4)
    action_table.style = 'Table Grid'
    
    # 表头
    headers = ["序号", "行动项", "负责人", "完成时间"]
    for i, header in enumerate(headers):
        action_table.rows[0].cells[i].text = header
        for paragraph in action_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()
    
    # 下次会议
    heading7 = doc.add_heading("下次会议", level=2)
    doc.add_paragraph("(请记录下次会议安排)")
    
    return doc


def create_contract(options: TemplateOptions) -> Document:
    """创建合同模板。"""
    doc = Document()
    
    # 标题
    title = doc.add_heading("合同", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 合同编号
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_para.add_run(f"合同编号: _____________")
    
    doc.add_paragraph()
    
    # 甲方乙方
    party_para = doc.add_paragraph()
    party_para.add_run("甲方（委托方）: ").bold = True
    party_para.add_run("________________________")
    doc.add_paragraph()
    
    party_para2 = doc.add_paragraph()
    party_para2.add_run("乙方（受托方）: ").bold = True
    party_para2.add_run("________________________")
    
    doc.add_paragraph()
    
    # 鉴于条款
    recitals = doc.add_paragraph()
    recitals.add_run("鉴于：").bold = True
    doc.add_paragraph("甲乙双方本着平等自愿、诚实信用的原则，经友好协商，就相关事宜达成如下协议：")
    
    # 合同条款
    doc.add_heading("第一条", level=2)
    doc.add_paragraph("(请在此填写第一条内容)")
    
    doc.add_heading("第二条", level=2)
    doc.add_paragraph("(请在此填写第二条内容)")
    
    doc.add_heading("第三条", level=2)
    doc.add_paragraph("(请在此填写第三条内容)")
    
    # 争议解决
    doc.add_heading("争议解决", level=2)
    doc.add_paragraph("本合同适用中华人民共和国法律。如发生争议，双方应友好协商解决；协商不成的，提交甲方所在地人民法院管辖。")
    
    doc.add_paragraph()
    
    # 签署
    doc.add_paragraph()
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sign_data = [
        ("甲方签字: _____________\n\n日期: ____年____月____日",
         "乙方签字: _____________\n\n日期: ____年____月____日"),
    ]
    
    for i, (left, right) in enumerate(sign_data):
        sign_table.rows[i].cells[0].text = left
        sign_table.rows[i].cells[1].text = right
    
    return doc


def create_resume(options: TemplateOptions) -> Document:
    """创建简历模板。"""
    colors = COLOR_SCHEMES.get(options.color_scheme, COLOR_SCHEMES["business"])
    
    doc = Document()
    
    # 姓名标题
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(options.title or "姓名")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = colors["primary"]
    
    # 联系信息
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = "电话: ___________ | 邮箱: ___________ | 地址: ___________"
    contact_run = contact_para.add_run(contact_info)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = colors["secondary"]
    
    # 分隔线
    doc.add_paragraph()
    
    # 个人简介
    doc.add_heading("个人简介", level=2)
    doc.add_paragraph("(请在此填写个人简介内容)")
    
    # 工作经历
    doc.add_heading("工作经历", level=2)
    
    exp_table = doc.add_table(rows=3, cols=3)
    exp_table.style = 'Table Grid'
    
    exp_headers = ["公司名称", "职位", "时间"]
    for i, header in enumerate(exp_headers):
        exp_table.rows[0].cells[i].text = header
        for paragraph in exp_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 教育背景
    doc.add_heading("教育背景", level=2)
    doc.add_paragraph("(请在此填写教育背景)")
    
    # 技能特长
    doc.add_heading("技能特长", level=2)
    doc.add_paragraph("(请在此填写技能特长)")
    
    # 证书荣誉
    doc.add_heading("证书荣誉", level=2)
    doc.add_paragraph("(请在此填写证书和荣誉)")
    
    return doc


# 模板注册表
TEMPLATES = {
    "academic_paper": create_academic_paper,
    "business_report": create_business_report,
    "meeting_minutes": create_meeting_minutes,
    "contract": create_contract,
    "resume": create_resume,
}


def create_from_template(
    template_name: str,
    options: TemplateOptions
) -> Document | None:
    """从模板创建文档。
    
    Args:
        template_name: 模板名称
        options: 模板选项
        
    Returns:
        Document 对象，失败返回 None
    """
    creator = TEMPLATES.get(template_name)
    if creator:
        return creator(options)
    return None


def list_templates() -> list[dict[str, str]]:
    """列出所有可用模板。"""
    return [
        {"name": "academic_paper", "title": "学术论文", "description": "包含封面、目录、摘要的学术论文模板"},
        {"name": "business_report", "title": "商业报告", "description": "包含执行摘要、分析建议的商业报告模板"},
        {"name": "meeting_minutes", "title": "会议纪要", "description": "包含议程、决议、行动项的会议纪要模板"},
        {"name": "contract", "title": "合同协议", "description": "包含标准条款的合同模板"},
        {"name": "resume", "title": "个人简历", "description": "专业简历模板"},
    ]


# 用于测试
if __name__ == "__main__":
    from datetime import datetime
    
    options = TemplateOptions(
        title="测试文档",
        author="张三",
        date=datetime.now().strftime("%Y年%m月%d日"),
        organization="测试公司",
        color_scheme="business"
    )
    
    # 测试创建商业报告
    doc = create_from_template("business_report", options)
    if doc:
        output_path = "test_template.docx"
        doc.save(output_path)
        print(f"已创建: {output_path}")
    
    # 列出模板
    print("\n可用模板:")
    for tmpl in list_templates():
        print(f"  - {tmpl['title']}: {tmpl['description']}")
